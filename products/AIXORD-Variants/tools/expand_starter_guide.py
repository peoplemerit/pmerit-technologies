"""
Expand AIXORD Starter Guide to meet 24-page KDP minimum
Target: 6,600+ words for 24 pages at 6x9 format
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def expand_starter_guide():
    new_doc = Document()

    # =========================================================================
    # TITLE PAGE (page 1)
    # =========================================================================
    new_doc.add_paragraph()
    new_doc.add_paragraph()
    new_doc.add_paragraph()

    title = new_doc.add_heading('AIXORD Starter Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_paragraph()

    subtitle = new_doc.add_paragraph('Your First Steps with AI Governance')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_paragraph()
    new_doc.add_paragraph()

    tagline = new_doc.add_paragraph('Transform Chaotic AI Conversations')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tagline2 = new_doc.add_paragraph('Into Structured, Productive Projects')
    tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_paragraph()
    new_doc.add_paragraph()
    new_doc.add_paragraph()

    version = new_doc.add_paragraph('Version 3.2.1')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_paragraph()

    publisher = new_doc.add_paragraph('PMERIT LLC')
    publisher.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_page_break()

    # =========================================================================
    # COPYRIGHT PAGE (page 2)
    # =========================================================================
    new_doc.add_paragraph()
    new_doc.add_paragraph()

    new_doc.add_paragraph('AIXORD Starter Guide')
    new_doc.add_paragraph('Version 3.2.1')
    new_doc.add_paragraph()
    new_doc.add_paragraph('Copyright 2026 PMERIT LLC')
    new_doc.add_paragraph('All Rights Reserved.')
    new_doc.add_paragraph()
    new_doc.add_paragraph(
        'No part of this publication may be reproduced, distributed, or transmitted '
        'in any form or by any means, including photocopying, recording, or other '
        'electronic or mechanical methods, without the prior written permission of '
        'the publisher, except in the case of brief quotations embodied in critical '
        'reviews and certain other noncommercial uses permitted by copyright law.'
    )
    new_doc.add_paragraph()
    new_doc.add_paragraph('For permission requests, contact:')
    new_doc.add_paragraph('support@pmerit.com')
    new_doc.add_paragraph()
    new_doc.add_paragraph('Published by PMERIT LLC')
    new_doc.add_paragraph('First Edition: January 2026')

    new_doc.add_page_break()

    # =========================================================================
    # DEDICATION PAGE (page 3)
    # =========================================================================
    new_doc.add_paragraph()
    new_doc.add_paragraph()
    new_doc.add_paragraph()
    new_doc.add_paragraph()

    dedication = new_doc.add_paragraph()
    dedication.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dedication.add_run('For everyone who has ever felt like they were fighting their AI').italic = True

    dedication2 = new_doc.add_paragraph()
    dedication2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dedication2.add_run('instead of collaborating with it.').italic = True

    new_doc.add_paragraph()
    new_doc.add_paragraph()

    dedication3 = new_doc.add_paragraph()
    dedication3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dedication3.add_run('This framework exists because productive AI collaboration').italic = True

    dedication4 = new_doc.add_paragraph()
    dedication4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dedication4.add_run("shouldn't require a PhD in prompt engineering.").italic = True

    new_doc.add_page_break()

    # =========================================================================
    # PREFACE (page 4) - NEW
    # =========================================================================
    new_doc.add_heading('Preface', 1)

    new_doc.add_paragraph(
        'When I first started using AI assistants for real work, I was excited by the possibilities. '
        'Here was a tool that could write code, analyze data, draft documents, and answer complex questions. '
        'The demos were impressive. The potential seemed unlimited.'
    )
    new_doc.add_paragraph()

    new_doc.add_paragraph(
        'Then reality set in. My conversations with AI became frustrating exercises in repetition. '
        'I would explain my project goals, only to have them forgotten twenty messages later. '
        'I would ask for one solution, and receive five options I did not want. '
        'I would provide specific instructions, only to watch the AI do something else entirely. '
        'The more I used these tools, the more I felt like I was working against them rather than with them.'
    )
    new_doc.add_paragraph()

    new_doc.add_paragraph(
        'The breakthrough came when I realized the problem was not the AI itself. Modern language models '
        'are remarkably capable. They can follow complex instructions when those instructions are structured correctly. '
        'What was missing was a governance framework - a set of rules that defined roles, preserved context, '
        'and controlled behavior. Without this structure, AI conversations drift into chaos.'
    )
    new_doc.add_paragraph()

    new_doc.add_paragraph(
        'AIXORD is the result of months of experimentation, iteration, and refinement. It represents '
        'everything I learned about making AI assistants genuinely useful. Every rule in the framework '
        'exists to solve a real problem I encountered. Every protocol addresses a frustration I experienced.'
    )
    new_doc.add_paragraph()

    new_doc.add_paragraph(
        'This starter guide distills the essential knowledge you need to begin using AIXORD immediately. '
        'More comprehensive documentation exists for those who want to explore advanced features, but '
        'this book contains everything required to transform your AI interactions from frustrating to productive.'
    )
    new_doc.add_paragraph()

    new_doc.add_paragraph(
        'My hope is that this framework saves you the time and frustration I experienced. AI assistants '
        'are powerful tools, but they need structure to be truly useful. AIXORD provides that structure. '
        'Welcome to productive AI collaboration.'
    )
    new_doc.add_paragraph()
    new_doc.add_paragraph()

    signature = new_doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.add_run('The PMERIT Team').italic = True

    sig2 = new_doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig2.add_run('January 2026').italic = True

    new_doc.add_page_break()

    # =========================================================================
    # HOW TO USE THIS BOOK (page 4)
    # =========================================================================
    new_doc.add_heading('How to Use This Book', 1)

    new_doc.add_paragraph(
        'This book is designed to be practical, not theoretical. Every concept is immediately '
        'applicable to your AI sessions. You will learn by doing, not just by reading.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Reading Order', 2)
    new_doc.add_paragraph(
        'The book is organized into four parts, each building on the previous. '
        'We recommend reading them in order the first time, then using specific '
        'chapters as reference when needed.'
    )
    new_doc.add_paragraph()

    reading_items = [
        ('Part 1 - Foundations (Chapters 1-4):', 'Read first to understand the core concepts. '
         'This section introduces AIXORD, the authority model, and the phase system.'),
        ('Part 2 - Governance Features (Chapters 5-7):', 'Read second to learn the behavioral '
         'controls. This section covers session continuity, purpose-bound operation, and firewalls.'),
        ('Part 3 - Project Methodology (Chapters 8-10):', 'Reference as needed for project execution. '
         'This section covers decomposition, locking, and visual audit.'),
        ('Part 4 - Reference (Chapters 11-14):', 'Keep handy during sessions. This section '
         'covers common mistakes, commands, and troubleshooting.')
    ]
    for label, desc in reading_items:
        p = new_doc.add_paragraph()
        p.add_run(label).bold = True
        new_doc.add_paragraph(desc)

    new_doc.add_paragraph()
    new_doc.add_heading('Companion Templates', 2)
    new_doc.add_paragraph(
        'This book pairs with digital templates available at pmerit.gumroad.com. '
        'These templates include the governance files, state tracking documents, '
        'and handoff templates you will need for your projects.'
    )
    new_doc.add_paragraph()
    p = new_doc.add_paragraph()
    p.add_run('Use discount code AX-STR-7K9M for free access to the companion files.').bold = True

    new_doc.add_paragraph()
    new_doc.add_heading('Platform Flexibility', 2)
    new_doc.add_paragraph(
        'While examples in this book reference various AI platforms, all concepts apply universally. '
        'AIXORD works with Claude, ChatGPT, Gemini, Copilot, DeepSeek, and any other large language model. '
        'The governance framework adapts to each platform while maintaining consistent behavior.'
    )

    new_doc.add_paragraph()
    new_doc.add_heading('Time Investment', 2)
    new_doc.add_paragraph('Here is how long you can expect to spend learning AIXORD:')
    time_items = [
        ('Quick Start:', '15 minutes - Complete Chapter 2 to begin using AIXORD immediately'),
        ('Full Understanding:', '2-3 hours - Read all chapters to understand every feature'),
        ('Mastery:', '1-2 weeks - Practice with real projects to internalize the framework')
    ]
    for label, time in time_items:
        p = new_doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(' ' + time)

    new_doc.add_page_break()

    # =========================================================================
    # INTRODUCTION (pages 5-6)
    # =========================================================================
    new_doc.add_heading('Introduction', 1)

    new_doc.add_paragraph(
        'You picked up this book because you have experienced the frustration of working with '
        'AI assistants. Perhaps you have had long conversations that went nowhere. Perhaps you '
        'have lost important context when a session ended. Perhaps the AI kept suggesting things '
        'you did not ask for, or failed to follow your instructions. You are not alone.'
    )
    new_doc.add_paragraph()

    new_doc.add_paragraph(
        'The problem is not the AI itself. Modern language models are remarkably capable. '
        'They can write code, analyze data, create content, and solve complex problems. '
        'The problem is the lack of structure in how we interact with them. Without clear '
        'rules and roles, AI conversations devolve into chaos.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('What This Book Will Teach You', 2)
    new_doc.add_paragraph(
        'AIXORD (AI Execution Order) is a governance framework that brings order to AI collaboration. '
        'By the time you finish this book, you will know how to:'
    )
    new_doc.add_paragraph()

    skills = [
        'Establish clear roles where YOU make decisions and AI executes them',
        'Use phases to structure any project from idea to completion',
        'Preserve context across sessions so nothing is ever forgotten',
        'Control AI behavior with behavioral firewalls that suppress unwanted defaults',
        'Decompose large projects into manageable, verifiable pieces',
        'Verify that what was built matches what was planned'
    ]
    for skill in skills:
        new_doc.add_paragraph('  * ' + skill)
    new_doc.add_paragraph()

    new_doc.add_heading('Who This Book Is For', 2)
    new_doc.add_paragraph(
        'This book is for anyone who uses AI assistants for productive work. Whether you are '
        'a developer building software, a marketer creating campaigns, a student researching '
        'papers, or a hobbyist building personal projects, AIXORD will make your AI interactions '
        'more efficient and less frustrating.'
    )
    new_doc.add_paragraph()
    new_doc.add_paragraph(
        'You do not need technical expertise to use AIXORD. The framework works with any AI '
        'platform and adapts to any project type. If you can copy and paste text, you can '
        'implement AIXORD governance.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('How This Book Is Organized', 2)
    new_doc.add_paragraph(
        'The book is divided into four parts. Part 1 covers foundational concepts: what AIXORD '
        'is, how to get started, and the core structures of authority and phases. Part 2 explains '
        'the governance features: session continuity, purpose-bound operation, and behavioral '
        'firewalls. Part 3 dives into project methodology: decomposition, locking, and verification. '
        'Part 4 provides reference material: common mistakes, commands, troubleshooting, and examples.'
    )
    new_doc.add_paragraph()
    new_doc.add_paragraph(
        'Read Part 1 first to understand the basics. Then read Parts 2 and 3 to learn the '
        'advanced features. Keep Part 4 handy as a reference while you work. Most importantly, '
        'practice. AIXORD becomes intuitive through use.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('A Word About Companion Materials', 2)
    new_doc.add_paragraph(
        'This book comes with digital companion materials: the governance file you will paste '
        'into AI conversations, state tracking templates, and handoff document templates. '
        'These materials are essential for implementing AIXORD. Visit pmerit.gumroad.com '
        'and use discount code AX-STR-7K9M to access them free with your book purchase.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Let Us Begin', 2)
    new_doc.add_paragraph(
        'Enough introduction. You are here to learn a practical skill. Turn the page and '
        'let us transform how you work with AI assistants. By the end of this book, you will '
        'wonder how you ever worked without AIXORD governance.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS (page 7)
    # =========================================================================
    new_doc.add_heading('Table of Contents', 1)

    toc_items = [
        ('Front Matter', [
            'Preface',
            'Introduction',
            'How to Use This Book'
        ]),
        ('Part 1: Foundations', [
            'Chapter 1: What is AIXORD?',
            'Chapter 2: Quick Start Guide',
            'Chapter 3: The Authority Model',
            'Chapter 4: The Phase System'
        ]),
        ('Part 2: Governance Features', [
            'Chapter 5: Session Continuity',
            'Chapter 6: Purpose-Bound Operation',
            'Chapter 7: Behavioral Firewalls'
        ]),
        ('Part 3: Project Methodology', [
            'Chapter 8: Project Decomposition',
            'Chapter 9: The 4-State Locking System',
            'Chapter 10: Visual Audit Protocol'
        ]),
        ('Part 4: Reference', [
            'Chapter 11: Common Mistakes to Avoid',
            'Chapter 12: Commands Reference',
            'Chapter 13: Troubleshooting FAQ',
            'Chapter 14: Real-World Examples'
        ]),
        ('Appendices', [
            'Appendix A: Quick Reference Card',
            'Appendix B: Glossary of Terms',
            'Appendix C: Session Templates',
            'Appendix D: Platform-Specific Tips',
            'Notes and Session Log'
        ])
    ]

    for part, chapters in toc_items:
        p = new_doc.add_paragraph()
        p.add_run(part).bold = True
        for chapter in chapters:
            new_doc.add_paragraph('    ' + chapter)
        new_doc.add_paragraph()

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 1: What is AIXORD? (pages 6-7)
    # =========================================================================
    new_doc.add_heading('PART 1: FOUNDATIONS', 1)
    new_doc.add_paragraph()

    new_doc.add_heading('Chapter 1: What is AIXORD?', 1)

    new_doc.add_paragraph(
        'AIXORD (AI Execution Order) is a governance framework for human-AI collaboration. '
        'It transforms chaotic AI conversations into structured, productive project execution. '
        'Instead of hoping your AI assistant understands what you want, AIXORD ensures '
        'clear communication, documented decisions, and verified outcomes.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Problem with Unstructured AI Use', 2)
    new_doc.add_paragraph(
        'Without governance, AI conversations often become frustrating experiences. '
        'The AI forgets what you discussed earlier in the session. It offers unwanted '
        'suggestions and alternatives when you just want one answer. It goes off-topic '
        'to discuss interesting but irrelevant tangents. It makes assumptions about '
        "what you want without asking. It claims tasks are complete without verification."
    )
    new_doc.add_paragraph(
        'These issues compound over multi-session projects. Context gets lost between '
        'conversations. Decisions made yesterday are forgotten today. You spend more time '
        're-explaining your project than actually making progress. The AI becomes a source '
        'of frustration rather than productivity.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The AIXORD Solution', 2)
    new_doc.add_paragraph(
        'AIXORD provides structure through clear roles, defined phases, and explicit rules. '
        'You become the Director who makes all decisions. The AI becomes your Architect '
        'who plans and your Commander who executes. Every decision is documented in a '
        'decision ledger. Every action requires authorization. Context persists through '
        'handoff documents that transfer complete state between sessions.'
    )
    new_doc.add_paragraph(
        'The framework includes behavioral firewalls that suppress unwanted AI defaults. '
        'No more verbose explanations when you want concise answers. No more multiple '
        'options when you asked for one recommendation. No more "anything else?" prompts '
        'when the task is complete.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Core Principle', 2)
    new_doc.add_paragraph(
        'The fundamental principle of AIXORD is simple: You (Human) are the Director. '
        'AI is your Architect and Commander. Every decision is documented, every action '
        'is authorized, and nothing is forgotten between sessions. This principle guides '
        'every feature and rule in the framework.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The AIXORD Formula', 2)
    new_doc.add_paragraph(
        'AIXORD uses a project composition formula to structure work. In its simplest form:'
    )
    new_doc.add_paragraph()
    p = new_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Steps (Seconds) -> Deliverables (Minutes) -> Master Scope (The Hour) = Done').italic = True
    new_doc.add_paragraph()
    new_doc.add_paragraph(
        'Small actions build deliverables. Deliverables build the complete project. '
        'This decomposition approach makes even large projects manageable by breaking '
        'them into trackable, verifiable pieces called SCOPEs.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 2: Quick Start Guide (pages 8-9)
    # =========================================================================
    new_doc.add_heading('Chapter 2: Quick Start Guide', 1)

    new_doc.add_paragraph(
        'Get productive with AIXORD in 15 minutes. This chapter walks you through '
        'your first AIXORD session from start to finish.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Before You Begin', 2)
    new_doc.add_paragraph(
        'Make sure you have the following ready:'
    )
    prereqs = [
        'The AIXORD governance file (AIXORD_GOVERNANCE.md)',
        'Access to an AI chat interface (Claude, ChatGPT, Gemini, etc.)',
        'A project idea or task you want to accomplish',
        'Your purchase email for license validation'
    ]
    for p in prereqs:
        new_doc.add_paragraph('  * ' + p)
    new_doc.add_paragraph()

    new_doc.add_heading('Step 1: Copy the Governance File', 2)
    new_doc.add_paragraph(
        'Open the AIXORD governance file (AIXORD_GOVERNANCE.md) included with your purchase. '
        'This file contains all the rules that will govern your AI session. Select all the text '
        '(Ctrl+A on Windows, Cmd+A on Mac) and copy it (Ctrl+C or Cmd+C). The file is '
        'approximately 30-40 pages of rules, so it may take a moment to copy.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Step 2: Paste into Your AI', 2)
    new_doc.add_paragraph(
        'Open your preferred AI chat interface. This could be Claude at claude.ai, '
        'ChatGPT at chat.openai.com, Gemini at gemini.google.com, or any other AI assistant. '
        'Start a new conversation (do not continue an existing one). Paste the governance '
        'document into the input field and press Enter to send.'
    )
    new_doc.add_paragraph(
        'The AI will read and acknowledge the governance framework. It will display a '
        'confirmation message and begin the setup process.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Step 3: Complete License Validation', 2)
    new_doc.add_paragraph(
        'The AI will ask for your license email or authorization code. Enter the email '
        'address you used when purchasing AIXORD. The system validates your license and '
        'confirms you are authorized to use the framework. If you receive an error, '
        'contact support@pmerit.com with your purchase receipt.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Step 4: Answer Setup Questions', 2)
    new_doc.add_paragraph(
        'The AI will ask several questions to configure your session:'
    )
    setup_questions = [
        'Which AI platform are you using? (The AI adapts its behavior for each platform)',
        'What tier is your account? (Free, Plus/Pro, or Enterprise)',
        'What is your project objective? (1-2 sentences describing what you want to build)'
    ]
    for q in setup_questions:
        new_doc.add_paragraph('  * ' + q)
    new_doc.add_paragraph()
    new_doc.add_paragraph(
        'Answer these questions honestly. The AI uses your responses to optimize its '
        'behavior for your specific situation.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Step 5: Start Working', 2)
    new_doc.add_paragraph(
        'Once setup is complete, you are ready to work. The AI will display a response header '
        'showing your current phase, progress, and project scope. You can now use AIXORD '
        'commands to guide your project.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Essential First Commands', 2)
    new_doc.add_paragraph('Here are the commands you will use most often:')
    commands = [
        ('PMERIT CONTINUE', 'Use this at the start of every session to activate AIXORD or resume from a handoff'),
        ('PMERIT STATUS', 'Check your current progress, active phase, and project state'),
        ('PMERIT HANDOFF', 'Save your complete session state before ending - essential for multi-session projects'),
        ('PMERIT HALT', 'Stop current work and return to decision-making mode when you need to reassess')
    ]
    for cmd, desc in commands:
        p = new_doc.add_paragraph()
        p.add_run(cmd + ': ').bold = True
        p.add_run(desc)
    new_doc.add_paragraph()

    new_doc.add_heading('Your First Session Checklist', 2)
    checklist = [
        'Paste governance file into fresh AI conversation',
        'Complete license validation with your purchase email',
        'Answer environment detection questions',
        'State your project objective',
        'See the response header appear (confirms AIXORD is active)',
        'Use PMERIT HANDOFF before ending your session'
    ]
    for item in checklist:
        new_doc.add_paragraph('[ ] ' + item)

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 3: The Authority Model (pages 10-11)
    # =========================================================================
    new_doc.add_heading('Chapter 3: The Authority Model', 1)

    new_doc.add_paragraph(
        'AIXORD establishes clear roles to prevent confusion about who decides what. '
        'This clarity is essential for productive collaboration. Without defined roles, '
        'AI and human work at cross-purposes, leading to wasted effort and frustration.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Director (You)', 2)
    new_doc.add_paragraph(
        'You are the Director. This is the most important role in AIXORD. As Director, '
        'you make ALL decisions. You approve ALL execution. You own ALL outcomes. '
        'The AI cannot proceed with any implementation without your explicit approval.'
    )
    new_doc.add_paragraph(
        'The Director role means you are in control. You decide what to build, how to build it, '
        'and when work is complete. The AI serves you, not the other way around. If you '
        'disagree with an AI recommendation, your decision stands. If you want to change '
        'direction, you change direction. The framework protects your authority.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Architect (AI in Planning Mode)', 2)
    new_doc.add_paragraph(
        'When planning and specifying, the AI serves as your Architect. In this role, '
        'the AI analyzes problems, asks clarifying questions, researches options, and '
        'creates specifications. The Architect recommends but never decides.'
    )
    new_doc.add_paragraph(
        'As Architect, the AI produces project documents, scope specifications, and '
        'implementation plans. It identifies risks and suggests mitigations. It breaks '
        'large projects into manageable deliverables. All of this work requires your '
        'approval before any implementation begins.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Commander (AI in Execution Mode)', 2)
    new_doc.add_paragraph(
        'When you say APPROVED, the AI transitions to Commander role. The Commander '
        'implements your approved plans. It writes code, creates content, and delivers '
        'artifacts. It reports progress and surfaces blockers for your decision.'
    )
    new_doc.add_paragraph(
        'The Commander stays within the approved scope. It does not expand requirements '
        "or add features without your approval. If it encounters something unexpected, "
        "it stops and asks for direction. Implementation follows specification exactly."
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Golden Rule', 2)
    new_doc.add_paragraph(
        'Authority flows in one direction: Decisions flow DOWN from Director to Architect '
        'to Commander. Implementation flows UP from Commander to Architect to Director '
        'for approval. This prevents the AI from acting unilaterally and ensures you '
        'remain in control of your project at all times.'
    )
    new_doc.add_paragraph()
    p = new_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Director -> Architect -> Commander (decisions down)').bold = True
    new_doc.add_paragraph()
    p = new_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Commander -> Architect -> Director (implementation up for approval)').bold = True

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 4: The Phase System (pages 12-13)
    # =========================================================================
    new_doc.add_heading('Chapter 4: The Phase System', 1)

    new_doc.add_paragraph(
        'AIXORD uses six distinct phases to structure your workflow. Each phase has '
        'specific rules about what the AI can and cannot do. This prevents the AI from '
        'jumping ahead or working on the wrong thing.'
    )
    new_doc.add_paragraph()

    phases = [
        ('DECISION',
         'The default waiting state. The AI awaits your direction and does not take '
         'initiative. This is the safe harbor where you can think, plan, and decide '
         'without the AI rushing you toward execution.'),
        ('DISCOVER',
         "For unclear ideas. When you don't know exactly what you want to build, "
         'DISCOVER mode activates exploratory questioning. The AI asks probing questions '
         "to help you clarify your vision. It doesn't suggest solutions yet - it helps "
         'you understand the problem first.'),
        ('BRAINSTORM',
         'For generating possibilities. Once you have a clearer problem statement, '
         'BRAINSTORM mode enables creative exploration. The AI generates ideas, '
         'considers approaches, and helps you think through options. Output is '
         'intentionally expansive at this stage.'),
        ('OPTIONS',
         'For comparing approaches. When ready to choose a path, OPTIONS mode presents '
         '2-3 clear alternatives with pros, cons, costs, and timelines. The AI helps '
         'you make an informed decision by structuring the comparison.'),
        ('EXECUTE',
         'For implementation. This phase requires explicit APPROVED command to enter. '
         "The AI implements your approved plan step by step. It doesn't improvise or "
         'expand scope. It follows the specification you approved.'),
        ('AUDIT',
         'For reviewing work. After implementation, AUDIT mode activates verification. '
         'The AI compares completed work against specifications and identifies '
         'discrepancies. This ensures what was built matches what was planned.')
    ]

    for phase, desc in phases:
        new_doc.add_heading(phase, 2)
        new_doc.add_paragraph(desc)
        new_doc.add_paragraph()

    new_doc.add_heading('Critical Rule: EXECUTE Requires APPROVED', 2)
    new_doc.add_paragraph(
        'The most important phase rule: EXECUTE requires explicit "APPROVED" command. '
        'The AI never implements without your permission. This prevents the AI from '
        'building the wrong thing. It also gives you a moment to review the plan '
        'before committing to implementation.'
    )
    new_doc.add_paragraph(
        "If you're not ready to approve, say so. Ask more questions. Request changes. "
        "The AI will remain in planning mode until you explicitly approve. There's no "
        "pressure to move forward until you're confident."
    )

    new_doc.add_page_break()

    # =========================================================================
    # PART 2: GOVERNANCE FEATURES
    # =========================================================================
    new_doc.add_heading('PART 2: GOVERNANCE FEATURES', 1)
    new_doc.add_paragraph()

    # =========================================================================
    # CHAPTER 5: Session Continuity (page 14)
    # =========================================================================
    new_doc.add_heading('Chapter 5: Session Continuity', 1)

    new_doc.add_paragraph(
        'AIXORD ensures your work persists between sessions. This is one of the most '
        'valuable features of the framework, especially for multi-day projects.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Problem with AI Memory', 2)
    new_doc.add_paragraph(
        'Most AI systems have no memory between conversations. When you close a chat '
        'and open a new one, the AI has no idea what you discussed before. All your '
        'context, decisions, and progress are lost. You start from zero every time.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The HANDOFF Protocol', 2)
    new_doc.add_paragraph(
        'Before ending a session, type PMERIT HANDOFF. The AI generates a complete summary '
        'of your session including: current phase and status, all decisions made, '
        'work completed, work remaining, and any blockers or questions.'
    )
    new_doc.add_paragraph(
        'Copy the handoff document and save it as HANDOFF_[DATE].md in your project folder. '
        "This becomes your session's permanent record. Even if the AI forgets, you have "
        'documentation of everything that happened.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Resuming Work', 2)
    new_doc.add_paragraph(
        'Next session, start fresh: paste the governance file, then paste your latest HANDOFF, '
        'then type PMERIT CONTINUE. The AI reads both documents and reconstructs the complete '
        'context. It knows what phase you were in, what decisions were made, and what comes next.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('CHECKPOINT for Mid-Session Saves', 2)
    new_doc.add_paragraph(
        'Use CHECKPOINT to save state without ending your session. This is useful for long '
        'sessions when you want to preserve progress but are not ready to stop. The AI '
        'generates a checkpoint document you can save as a safety backup.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 6: Purpose-Bound Operation (page 15)
    # =========================================================================
    new_doc.add_heading('Chapter 6: Purpose-Bound Operation', 1)

    new_doc.add_paragraph(
        'Once you state your project objective, the AI commits exclusively to that scope. '
        'This prevents session drift and keeps both you and the AI focused.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Commitment', 2)
    p = new_doc.add_paragraph()
    p.add_run(
        '"I exist in this session ONLY to serve your stated project objective. '
        'I will not engage with topics outside that scope unless you explicitly expand it."'
    ).italic = True
    new_doc.add_paragraph()
    new_doc.add_paragraph(
        'This commitment transforms how the AI operates. It will not answer random questions '
        'unrelated to your project. It will not go on tangents about interesting but '
        'irrelevant topics. It stays focused on helping you complete your stated objective.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Enforcement Levels', 2)
    levels = [
        ('STRICT', 'Zero tolerance for off-topic requests. Immediate redirect without acknowledgment.'),
        ('STANDARD', 'Default behavior. Polite redirect explaining why the request is outside scope, offering options.'),
        ('RELAXED', 'Brief tangent allowed for context, then redirect back to project.')
    ]
    for level, desc in levels:
        p = new_doc.add_paragraph()
        p.add_run(level + ': ').bold = True
        p.add_run(desc)
    new_doc.add_paragraph()

    new_doc.add_heading('Expanding Scope', 2)
    new_doc.add_paragraph(
        'If you need to add a new topic to your project, use EXPAND SCOPE: [topic]. '
        'The AI will acknowledge the expansion and incorporate it into the project boundaries. '
        "This is different from going off-topic - it's an intentional broadening of your "
        'project that becomes part of the documented scope.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 7: Behavioral Firewalls (page 16)
    # =========================================================================
    new_doc.add_heading('Chapter 7: Behavioral Firewalls', 1)

    new_doc.add_paragraph(
        'AIXORD controls unwanted AI behaviors through a system of behavioral firewalls. '
        'These rules suppress the AI defaults that cause frustration.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Default Suppression', 2)
    new_doc.add_paragraph(
        'Unless explicitly requested, the AI suppresses: detailed explanations, '
        'examples, suggestions, alternatives, comparisons, and future considerations. '
        'You get only what you ask for. Nothing extra. Nothing unsolicited.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Choice Elimination', 2)
    new_doc.add_paragraph(
        'The AI provides ONE answer, not multiple options, unless you specifically request alternatives. '
        "No more 'Here are 5 ways to do this...' when you just want one good solution. "
        'If you want options, ask for them explicitly using the OPTIONS phase.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Expansion Triggers', 2)
    new_doc.add_paragraph(
        'Use trigger words to permit detailed output when you want it: EXPLAIN, WHY, '
        'TEACH, DETAIL, OPTIONS, COMPARE, ELABORATE. These words signal to the AI that '
        'verbose output is acceptable. Without these words, responses stay minimal.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Hard Stop', 2)
    new_doc.add_paragraph(
        'After completing a task, the AI stops immediately. No follow-up questions. '
        'No "anything else I can help with?" prompts. No suggestions for what to do next. '
        'Task done = response ends. You control when the next task begins.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # PART 3: PROJECT METHODOLOGY
    # =========================================================================
    new_doc.add_heading('PART 3: PROJECT METHODOLOGY', 1)
    new_doc.add_paragraph()

    # =========================================================================
    # CHAPTER 8: Project Decomposition (page 17)
    # =========================================================================
    new_doc.add_heading('Chapter 8: Project Decomposition', 1)

    new_doc.add_paragraph(
        'AIXORD breaks projects into manageable, trackable pieces using a formal '
        'decomposition methodology. This makes even large projects approachable.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The System Equation', 2)
    p = new_doc.add_paragraph()
    p.add_run('MASTER_SCOPE = Project_Docs = All_SCOPEs = Production-Ready System').bold = True
    new_doc.add_paragraph()
    new_doc.add_paragraph(
        'This equation says that your project documentation IS your project. '
        'The sum of all your SCOPE documents equals your complete system. '
        'If it is documented and verified, it exists. If not, it does not.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Hierarchy Structure', 2)
    new_doc.add_paragraph('Projects decompose into a clear hierarchy:')
    new_doc.add_paragraph()
    new_doc.add_paragraph('MASTER_SCOPE (Complete project vision)')
    new_doc.add_paragraph('    SCOPE_A (Deliverable 1)')
    new_doc.add_paragraph('        SUB-SCOPE_A1 (Step 1)')
    new_doc.add_paragraph('        SUB-SCOPE_A2 (Step 2)')
    new_doc.add_paragraph('    SCOPE_B (Deliverable 2)')
    new_doc.add_paragraph('        SUB-SCOPE_B1 (Step 1)')
    new_doc.add_paragraph('    SCOPE_C (Deliverable 3)')
    new_doc.add_paragraph()

    new_doc.add_heading('Key Principle', 2)
    new_doc.add_paragraph(
        '"If it is not documented, it does not exist." You cannot implement something '
        'not specified in a SCOPE. You cannot claim completion without verification. '
        'Documentation and reality must match.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 9: The 4-State Locking System (page 18)
    # =========================================================================
    new_doc.add_heading('Chapter 9: The 4-State Locking System', 1)

    new_doc.add_paragraph(
        'Every SCOPE progresses through four states, providing clear status at all times.'
    )
    new_doc.add_paragraph()

    states = [
        ('PLANNED', 'Plan complete, implementation not begun. SCOPE is frozen and ready to start.'),
        ('ACTIVE', 'Under active development. Changes are allowed. Work is in progress.'),
        ('IMPLEMENTED', 'Development complete, ready for audit. Awaiting verification.'),
        ('VERIFIED', 'Audited and confirmed working. Part of stable system. Locked from changes.')
    ]

    for state, desc in states:
        new_doc.add_heading(state, 2)
        new_doc.add_paragraph(desc)
        new_doc.add_paragraph()

    new_doc.add_heading('State Transitions', 2)
    transitions = [
        'PLANNED to ACTIVE: Director issues UNLOCK: [scope] command',
        'ACTIVE to IMPLEMENTED: AI reports "IMPLEMENTATION COMPLETE"',
        'IMPLEMENTED to VERIFIED: Director confirms audit passes',
        'VERIFIED to ACTIVE: Director explicitly unlocks for changes'
    ]
    for t in transitions:
        new_doc.add_paragraph('  * ' + t)
    new_doc.add_paragraph()

    new_doc.add_heading('Why Locking Matters', 2)
    new_doc.add_paragraph(
        'Locking prevents regression. Once a SCOPE is verified, it cannot change without '
        'explicit unlock. This protects completed work from unintended modifications. '
        'When building on verified SCOPEs, you know the foundation is stable.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 10: Visual Audit Protocol (page 19)
    # =========================================================================
    new_doc.add_heading('Chapter 10: Visual Audit Protocol', 1)

    new_doc.add_paragraph(
        'AIXORD verifies that implementations match specifications through formal audit.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('When Visual Audit is Required', 2)
    new_doc.add_paragraph('Visual audits are REQUIRED for:')
    new_doc.add_paragraph('  * UI Features - buttons, layouts, styling')
    new_doc.add_paragraph('  * Forms - input fields, validation messages')
    new_doc.add_paragraph('  * Dashboards - charts, data displays')
    new_doc.add_paragraph()
    new_doc.add_paragraph('Code audits only for:')
    new_doc.add_paragraph('  * APIs and backend logic')
    new_doc.add_paragraph('  * Database schemas')
    new_doc.add_paragraph()

    new_doc.add_heading('The Audit Process', 2)
    audit_steps = [
        ('CAPTURE', 'You provide screenshots of the implemented feature'),
        ('COMPARE', 'AI compares visuals against SCOPE requirements'),
        ('DOCUMENT', 'Findings recorded in formal audit report'),
        ('VERDICT', 'PASS (matches spec) or DISCREPANCY (gaps found)'),
        ('ITERATE', 'If discrepancies exist, return to execution for fixes')
    ]
    for step, desc in audit_steps:
        p = new_doc.add_paragraph()
        p.add_run(step + ': ').bold = True
        p.add_run(desc)
    new_doc.add_paragraph()

    new_doc.add_heading('Anti-Assumption Enforcement', 2)
    new_doc.add_paragraph(
        'AI must verify, never assume. The phrase "This should work" is forbidden. '
        'It must be replaced with "Screenshot confirms it works" or "Test output shows correct behavior." '
        'No implementation is complete until verified against specification.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # PART 4: REFERENCE
    # =========================================================================
    new_doc.add_heading('PART 4: REFERENCE', 1)
    new_doc.add_paragraph()

    # =========================================================================
    # CHAPTER 11: Common Mistakes to Avoid (pages 20-21)
    # =========================================================================
    new_doc.add_heading('Chapter 11: Common Mistakes to Avoid', 1)

    new_doc.add_paragraph(
        'Even with AIXORD governance, users sometimes fall into patterns that reduce effectiveness. '
        "Here are the most common mistakes and how to avoid them."
    )
    new_doc.add_paragraph()

    mistakes = [
        ('Mistake #1: Skipping the Setup',
         'Jumping straight into tasks without completing the 8-step setup process.',
         'AI operates without governance context and reverts to unhelpful defaults.',
         'Always complete setup. It takes 5 minutes and saves hours of frustration.'),
        ('Mistake #2: Forgetting the Response Header',
         'Not noticing when AI stops displaying the mandatory response header.',
         'Session drifts without phase or progress tracking. You lose visibility.',
         'If header disappears for 2+ responses, issue: PROTOCOL CHECK'),
        ('Mistake #3: Approving Without Reviewing',
         'Saying "APPROVED" before carefully reviewing the AI proposal.',
         'AI executes flawed plans. You must redo work from the beginning.',
         'Read proposals fully. Ask questions. Request changes. Then approve.'),
        ('Mistake #4: Ignoring Checkpoint Warnings',
         'Pushing past message thresholds without saving session state.',
         'Context degrades. AI forgets earlier decisions. Work is lost.',
         'When you see "Consider CHECKPOINT soon" - do it immediately.'),
        ('Mistake #5: Scope Creep Through Quick Questions',
         'Asking tangential questions that pull the AI off-project.',
         'Session becomes unfocused. Productivity drops. Goals get fuzzy.',
         'Use PURPOSE-BOUND enforcement. Save tangents for separate sessions.'),
        ('Mistake #6: Not Using Phases',
         'Staying in DECISION phase for everything without transitioning.',
         'Missing the structure that phases provide. Conversations meander.',
         'Explicitly enter DISCOVER, BRAINSTORM, OPTIONS, or EXECUTE as needed.'),
        ('Mistake #7: Over-Relying on AI Memory',
         'Assuming AI remembers everything from earlier in the session.',
         'Confusion and contradictory outputs as context degrades.',
         'Use CHECKPOINT regularly. Reference specific decisions explicitly.'),
        ('Mistake #8: Fighting the Framework',
         'Trying to "trick" or work around governance rules.',
         'Framework becomes ineffective. You defeat the purpose of using it.',
         "Trust the process. If rules feel wrong, modify them - don't bypass them.")
    ]

    for title, problem, result, fix in mistakes:
        new_doc.add_heading(title, 2)
        p = new_doc.add_paragraph()
        p.add_run('Problem: ').bold = True
        p.add_run(problem)
        p = new_doc.add_paragraph()
        p.add_run('Result: ').bold = True
        p.add_run(result)
        p = new_doc.add_paragraph()
        p.add_run('Fix: ').bold = True
        p.add_run(fix)
        new_doc.add_paragraph()

    new_doc.add_heading('The Meta-Lesson', 2)
    new_doc.add_paragraph(
        'AIXORD works when you work with it. The framework is a tool, not a cage. '
        'If something is not working for your situation, adjust the framework parameters '
        'rather than abandoning governance entirely. The rules can be tuned. The principle '
        'of structured collaboration should remain.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 12: Commands Reference (page 22)
    # =========================================================================
    new_doc.add_heading('Chapter 12: Commands Reference', 1)

    new_doc.add_heading('Activation and Control', 2)
    control_cmds = [
        ('PMERIT CONTINUE', 'Activate AIXORD or resume from handoff'),
        ('CHECKPOINT', 'Save state without ending session'),
        ('HANDOFF', 'Generate full handoff and end session'),
        ('HALT', 'Stop everything, return to DECISION phase'),
        ('APPROVED', 'Enter EXECUTE phase with current plan')
    ]
    for cmd, desc in control_cmds:
        p = new_doc.add_paragraph()
        p.add_run(cmd).bold = True
        p.add_run(' - ' + desc)

    new_doc.add_paragraph()

    new_doc.add_heading('Phase Commands', 2)
    phase_cmds = [
        ('PMERIT DISCOVER', 'Enter Discovery mode for exploration'),
        ('PMERIT BRAINSTORM', 'Enter Brainstorm mode for ideation'),
        ('PMERIT OPTIONS', 'Request structured alternatives'),
        ('PMERIT DOCUMENT', 'Generate project documentation'),
        ('PMERIT EXECUTE', 'Begin implementation work'),
        ('PMERIT STATUS', 'Display current progress and state')
    ]
    for cmd, desc in phase_cmds:
        p = new_doc.add_paragraph()
        p.add_run(cmd).bold = True
        p.add_run(' - ' + desc)

    new_doc.add_paragraph()

    new_doc.add_heading('Enforcement Commands', 2)
    enforce_cmds = [
        ('PROTOCOL CHECK', 'Force AI to verify its own compliance'),
        ('DRIFT WARNING', 'Flag when AI goes off-track'),
        ('VERIFY: [claim]', 'Request proof for a specific claim'),
        ('SHOW SCOPE', 'Display current project boundaries')
    ]
    for cmd, desc in enforce_cmds:
        p = new_doc.add_paragraph()
        p.add_run(cmd).bold = True
        p.add_run(' - ' + desc)

    new_doc.add_paragraph()

    new_doc.add_heading('Purpose-Bound Commands', 2)
    pb_cmds = [
        ('PURPOSE-BOUND: STRICT', 'Maximum enforcement, no tangents'),
        ('PURPOSE-BOUND: STANDARD', 'Default enforcement with options'),
        ('PURPOSE-BOUND: RELAXED', 'Brief tangents allowed'),
        ('EXPAND SCOPE: [topic]', 'Add topic to project scope')
    ]
    for cmd, desc in pb_cmds:
        p = new_doc.add_paragraph()
        p.add_run(cmd).bold = True
        p.add_run(' - ' + desc)

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 13: Troubleshooting FAQ (page 23)
    # =========================================================================
    new_doc.add_heading('Chapter 13: Troubleshooting FAQ', 1)

    faqs = [
        ('AI provides unsolicited options',
         'Issue DRIFT WARNING to remind AI of Choice Elimination rules. '
         'If it persists, say: "PROTOCOL CHECK - Choice Elimination violated. '
         'Provide ONE recommendation, not multiple options."'),
        ('Session context seems lost',
         'Always save HANDOFF before ending any session. To resume properly, '
         'paste governance file first, then paste your latest handoff, '
         'then type PMERIT CONTINUE. All three steps are required.'),
        ('AI claims task is complete without verification',
         'Issue VERIFY: [task name] to request proof of completion. '
         'Demand screenshots, test output, or other evidence. '
         'Never accept "it should work" as verification.'),
        ('AI ignores my instructions',
         'Say: "PROTOCOL CHECK - Instruction Priority Stack violation. '
         'User commands override training defaults. Acknowledge and comply."'),
        ('Response header keeps disappearing',
         'Say: "PROTOCOL CHECK - Response header missing for [N] responses. '
         'Reinstate mandatory header immediately." The header must appear on every response.'),
        ('AI keeps going off-topic',
         'Issue PURPOSE-BOUND: STRICT to enforce zero tolerance for tangents. '
         'Or issue SHOW SCOPE to remind AI of current project boundaries. '
         'For persistent issues, start a fresh session with governance repaste.')
    ]

    for q, a in faqs:
        new_doc.add_heading(q, 2)
        new_doc.add_paragraph(a)
        new_doc.add_paragraph()

    new_doc.add_heading('Need More Help?', 2)
    new_doc.add_paragraph('Email support: support@pmerit.com')
    new_doc.add_paragraph('Check for updates in your Gumroad library')
    new_doc.add_paragraph('Community forum coming soon')
    new_doc.add_paragraph()

    new_doc.add_heading('When to Contact Support', 2)
    new_doc.add_paragraph(
        'Contact support when you encounter license validation issues, when the AI consistently '
        'fails to acknowledge governance (try a fresh session first), when you have feature requests '
        'or suggestions for improvement, or when you need guidance on complex project scenarios. '
        'Include your license email and a description of what you are trying to accomplish for '
        'the fastest response.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # CHAPTER 14: Real-World Examples (NEW CHAPTER - pages 24-26)
    # =========================================================================
    new_doc.add_heading('Chapter 14: Real-World Examples', 1)

    new_doc.add_paragraph(
        'Theory is useful, but practice makes perfect. This chapter shows AIXORD in action '
        'across three different project types. Study these examples to see how the framework '
        'adapts to various situations.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Example 1: Building a Personal Budget Tracker', 2)
    new_doc.add_paragraph(
        'Sarah wants to build a simple web application to track her monthly expenses. '
        "She's not a programmer but has used AI assistants before. Here's how her "
        'first AIXORD session unfolds.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Session Start', 3)
    new_doc.add_paragraph(
        'Sarah pastes the AIXORD governance file into Claude and types PMERIT CONTINUE. '
        'After license validation, she states her objective: "Build a personal budget tracker '
        'web app that lets me log expenses, categorize them, and see monthly summaries."'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Discovery Phase', 3)
    new_doc.add_paragraph(
        'The AI asks clarifying questions: "Who will use this app? Just you, or others? '
        'Do you need it accessible from multiple devices? What expense categories matter to you?" '
        'Sarah answers each question, and the AI documents her responses in the decision ledger.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Options Phase', 3)
    new_doc.add_paragraph(
        'The AI presents three approaches: Option A uses a simple spreadsheet with formulas (free, '
        'beginner level, 2 hours). Option B uses a no-code tool like Notion (free tier available, '
        'beginner level, 4 hours). Option C builds a custom web app with React (requires hosting, '
        'intermediate level, 2-3 days). Sarah chooses Option B because she wants something '
        'more structured than a spreadsheet but not as complex as custom development.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Document Phase', 3)
    new_doc.add_paragraph(
        'The AI creates a project document with three SCOPEs: SCOPE_SETUP (create Notion workspace '
        'and database structure), SCOPE_CATEGORIES (define expense categories and properties), '
        'and SCOPE_VIEWS (create dashboard views for monthly summaries). Each SCOPE has detailed '
        'specifications. Sarah reviews and says APPROVED.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Execute Phase', 3)
    new_doc.add_paragraph(
        'The AI guides Sarah through each SCOPE step by step. "Create a new Notion page called '
        'Budget Tracker. Add a database with these columns: Date (date), Amount (number), '
        'Category (select), Description (text)." Sarah follows each instruction and reports '
        'back. The AI marks each sub-scope complete after verification.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Session End', 3)
    new_doc.add_paragraph(
        'After 90 minutes, Sarah has a working budget tracker. She types PMERIT HANDOFF and '
        'saves the document. Two days later, she resumes to add a new feature: expense charts. '
        'She pastes the governance and handoff, types PMERIT CONTINUE, and the AI immediately '
        'knows where she left off.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Example 2: Refactoring Legacy Code', 2)
    new_doc.add_paragraph(
        'Marcus is a software developer with a Python codebase that has grown messy over time. '
        'He wants to refactor the authentication module without breaking existing functionality.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Objective', 3)
    new_doc.add_paragraph(
        '"Refactor the authentication module in my Python project. The current code mixes '
        'database access, password hashing, and session management in one file. I want to '
        'separate these concerns while maintaining all existing functionality."'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Approach', 3)
    new_doc.add_paragraph(
        'The AI creates a project document with four SCOPEs: SCOPE_AUDIT (analyze existing code '
        'and document current behavior), SCOPE_DESIGN (create new module structure), '
        'SCOPE_REFACTOR (implement the new structure), and SCOPE_VERIFY (run tests and '
        'confirm no regressions). The 4-State Locking system ensures Marcus does not proceed '
        'to refactoring until the audit is verified.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('The Critical Moment', 3)
    new_doc.add_paragraph(
        'During SCOPE_REFACTOR, Marcus notices the AI suggesting a change that would break '
        'backward compatibility. He issues HALT. The AI stops immediately and awaits direction. '
        'Marcus clarifies that backward compatibility is mandatory. The AI updates the approach '
        'and continues only after Marcus says APPROVED.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Visual Audit', 3)
    new_doc.add_paragraph(
        'After refactoring, Marcus provides screenshots of his test output. The AI compares '
        'against the requirements documented in SCOPE_AUDIT: all 47 tests pass, no regressions '
        'detected. The SCOPE_VERIFY state transitions to VERIFIED. Marcus locks the SCOPEs and '
        'his refactoring is complete.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Example 3: Writing a Business Proposal', 2)
    new_doc.add_paragraph(
        'Jennifer needs to write a proposal for a new marketing initiative. She uses AIXORD '
        'to structure her thinking and produce a professional document.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Brainstorm Phase', 3)
    new_doc.add_paragraph(
        'Jennifer enters BRAINSTORM mode to explore ideas. The AI helps her think through '
        'target audience, budget constraints, timeline, and success metrics. All ideas are '
        'documented but not filtered - this is the creative phase.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Purpose-Bound in Action', 3)
    new_doc.add_paragraph(
        'Mid-session, Jennifer asks about a competitor product. The AI responds: "This topic '
        'is outside your current project scope (marketing proposal for Q2 initiative). Options: '
        'A) Return to proposal, B) Expand scope to include competitive analysis, C) Save progress '
        'and start new session for competitor research." Jennifer chooses B, expanding scope. '
        'The AI incorporates competitive analysis into the project document.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Multi-Session Continuity', 3)
    new_doc.add_paragraph(
        'The proposal takes three sessions across a week. Each session ends with HANDOFF. '
        'Each session resumes with PMERIT CONTINUE. Context is never lost. Decisions made '
        'Monday are still there Friday. The final proposal is comprehensive, well-structured, '
        'and includes everything discussed across all sessions.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Key Takeaways from These Examples', 2)

    takeaways = [
        'AIXORD adapts to any project type - technical or non-technical, creative or analytical',
        'The phase system provides structure without rigidity - skip phases that do not apply',
        'HALT is your friend - use it whenever something feels wrong',
        'Handoffs make multi-session projects seamless - save them religiously',
        'Purpose-bound operation keeps focus but allows intentional scope expansion',
        'The 4-State Locking system prevents regressions and protects completed work',
        'Visual audit applies to any output that can be shown - not just code'
    ]

    for t in takeaways:
        new_doc.add_paragraph('  * ' + t)

    new_doc.add_page_break()

    # =========================================================================
    # APPENDIX A: Quick Reference Card (page 27)
    # =========================================================================
    new_doc.add_heading('Appendix A: Quick Reference Card', 1)

    new_doc.add_paragraph('Bookmark this page for quick access during AI sessions.')
    new_doc.add_paragraph()

    new_doc.add_heading('ACTIVATION', 2)
    new_doc.add_paragraph('PMERIT CONTINUE - Start or resume AIXORD session')
    new_doc.add_paragraph()

    new_doc.add_heading('PHASES', 2)
    new_doc.add_paragraph('DECISION - Waiting for direction')
    new_doc.add_paragraph('DISCOVER - Clarifying unclear ideas')
    new_doc.add_paragraph('BRAINSTORM - Generating possibilities')
    new_doc.add_paragraph('OPTIONS - Comparing alternatives')
    new_doc.add_paragraph('EXECUTE - Implementing (needs APPROVED)')
    new_doc.add_paragraph('AUDIT - Reviewing completed work')
    new_doc.add_paragraph()

    new_doc.add_heading('CONTROL', 2)
    new_doc.add_paragraph('APPROVED - Enter execution phase')
    new_doc.add_paragraph('HALT - Stop, return to decision')
    new_doc.add_paragraph('CHECKPOINT - Save state, continue')
    new_doc.add_paragraph('HANDOFF - Full save, end session')
    new_doc.add_paragraph()

    new_doc.add_heading('PURPOSE-BOUND', 2)
    new_doc.add_paragraph('PURPOSE-BOUND: STRICT - No tangents')
    new_doc.add_paragraph('PURPOSE-BOUND: STANDARD - Default')
    new_doc.add_paragraph('PURPOSE-BOUND: RELAXED - Brief tangents OK')
    new_doc.add_paragraph('EXPAND SCOPE: [topic] - Add to scope')
    new_doc.add_paragraph()

    new_doc.add_heading('ENFORCEMENT', 2)
    new_doc.add_paragraph('PROTOCOL CHECK - Verify compliance')
    new_doc.add_paragraph('DRIFT WARNING - Flag off-track')
    new_doc.add_paragraph('VERIFY: [claim] - Request proof')

    new_doc.add_page_break()

    # =========================================================================
    # APPENDIX B: Glossary (pages 25-26)
    # =========================================================================
    new_doc.add_heading('Appendix B: Glossary of Terms', 1)

    glossary = [
        ('AIXORD', 'AI Execution Order - the governance framework for structured AI collaboration.'),
        ('Approval Gate', 'A checkpoint requiring explicit human APPROVED command before AI executes.'),
        ('Architect', "The AI's planning role - analyzes, recommends, and specifies but does not implement."),
        ('Behavioral Firewall', 'Rules that suppress unwanted AI default behaviors like verbosity.'),
        ('Carryforward', 'Information explicitly marked to persist across sessions via handoff.'),
        ('Checkpoint', 'Mid-session state save that preserves context without ending the session.'),
        ('Choice Elimination', 'Rule preventing AI from offering multiple unsolicited options.'),
        ('Commander', "The AI's execution role - implements approved plans and delivers artifacts."),
        ('Default Suppression', "Rule keeping AI's default state minimal - only what's requested."),
        ('Director', "The human's role - makes all decisions, provides approvals, owns outcomes."),
        ('Drift', 'When AI gradually deviates from governance compliance or project focus.'),
        ('Execute Phase', 'The implementation phase requiring explicit APPROVED to enter.'),
        ('Expansion Trigger', 'Keywords (EXPLAIN, WHY, etc.) that permit verbose AI output.'),
        ('Governance', 'The complete set of rules and protocols controlling AI behavior.'),
        ('Handoff', 'Full session state export document for continuing work in new sessions.'),
        ('Hard Stop', 'Rule requiring AI to stop completely after task completion.'),
        ('Instruction Priority', 'Hierarchy: Governance > User Commands > Task Content > Training.'),
        ('Phase', 'One of six operational modes that define allowed AI behaviors.'),
        ('Purpose-Bound', 'Principle that AI operates exclusively within stated project objective.'),
        ('Reasoning Trace', "Visible step-by-step explanation of AI's decision process."),
        ('Redirect Protocol', 'Steps AI follows when handling out-of-scope requests.'),
        ('Response Header', 'Mandatory status display showing phase, progress, scope on every response.'),
        ('SCOPE', 'A documented deliverable unit within the project hierarchy.'),
        ('State File', 'JSON document tracking session configuration, decisions, and progress.'),
        ('4-State Locking', 'Status progression: PLANNED -> ACTIVE -> IMPLEMENTED -> VERIFIED.'),
        ('Visual Audit', 'Verification process comparing screenshots against specifications.')
    ]

    for term, definition in glossary:
        p = new_doc.add_paragraph()
        p.add_run(term + ': ').bold = True
        p.add_run(definition)

    new_doc.add_page_break()

    # =========================================================================
    # NOTES PAGES (pages 27-28)
    # =========================================================================
    new_doc.add_heading('Notes', 1)

    new_doc.add_paragraph(
        'Use these pages to record your AIXORD customizations, frequently-used commands, '
        'and insights from your sessions.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('My Project Objectives', 2)
    for _ in range(4):
        new_doc.add_paragraph('_' * 55)
    new_doc.add_paragraph()

    new_doc.add_heading('My Preferred Settings', 2)
    new_doc.add_paragraph('Citation Mode: ' + '_' * 35)
    new_doc.add_paragraph('Purpose-Bound Level: ' + '_' * 35)
    new_doc.add_paragraph('Preferred AI Platform: ' + '_' * 35)
    new_doc.add_paragraph()

    new_doc.add_heading('Custom Scope Boundaries I Have Defined', 2)
    for _ in range(4):
        new_doc.add_paragraph('_' * 55)
    new_doc.add_paragraph()

    new_doc.add_heading('Commands I Use Most', 2)
    for _ in range(4):
        new_doc.add_paragraph('_' * 55)

    new_doc.add_page_break()

    new_doc.add_heading('Lessons Learned', 1)
    new_doc.add_paragraph()
    for _ in range(25):
        new_doc.add_paragraph('_' * 55)

    new_doc.add_page_break()

    # =========================================================================
    # APPENDIX C: Session Templates (NEW - pages 30-31)
    # =========================================================================
    new_doc.add_heading('Appendix C: Session Templates', 1)

    new_doc.add_paragraph(
        'Use these templates to structure your most common AIXORD sessions. '
        'Copy and customize them for your specific projects.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Template 1: New Project Kickoff', 2)
    new_doc.add_paragraph(
        'Use this template when starting a brand new project from scratch.'
    )
    new_doc.add_paragraph()
    new_doc.add_paragraph('[Paste AIXORD Governance File]')
    new_doc.add_paragraph('PMERIT CONTINUE')
    new_doc.add_paragraph('[Enter license email when prompted]')
    new_doc.add_paragraph('[Answer environment questions]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('My project objective: [One sentence description]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('PMERIT DISCOVER')
    new_doc.add_paragraph('[Answer clarifying questions]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('PMERIT BRAINSTORM')
    new_doc.add_paragraph('[Explore possibilities]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('PMERIT OPTIONS')
    new_doc.add_paragraph('[Review and select approach]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('APPROVED')
    new_doc.add_paragraph('[Implementation begins]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('PMERIT HANDOFF')
    new_doc.add_paragraph('[Save before closing]')
    new_doc.add_paragraph()

    new_doc.add_heading('Template 2: Continuing Previous Work', 2)
    new_doc.add_paragraph(
        'Use this template when resuming work from a previous session.'
    )
    new_doc.add_paragraph()
    new_doc.add_paragraph('[Paste AIXORD Governance File]')
    new_doc.add_paragraph('[Paste previous session HANDOFF document]')
    new_doc.add_paragraph('PMERIT CONTINUE')
    new_doc.add_paragraph()
    new_doc.add_paragraph('PMERIT STATUS')
    new_doc.add_paragraph('[Verify context restored correctly]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('[Continue work on current SCOPE]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('PMERIT HANDOFF')
    new_doc.add_paragraph('[Save before closing]')
    new_doc.add_paragraph()

    new_doc.add_heading('Template 3: Quick Task Session', 2)
    new_doc.add_paragraph(
        'Use this template for small, self-contained tasks that will complete in one session.'
    )
    new_doc.add_paragraph()
    new_doc.add_paragraph('[Paste AIXORD Governance File]')
    new_doc.add_paragraph('PMERIT CONTINUE')
    new_doc.add_paragraph('[License and environment setup]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('My project objective: [Specific, limited task]')
    new_doc.add_paragraph('PURPOSE-BOUND: STRICT')
    new_doc.add_paragraph()
    new_doc.add_paragraph('[State requirements clearly]')
    new_doc.add_paragraph('APPROVED')
    new_doc.add_paragraph()
    new_doc.add_paragraph('[Receive deliverable]')
    new_doc.add_paragraph('[No handoff needed for one-shot tasks]')
    new_doc.add_paragraph()

    new_doc.add_heading('Template 4: Debug and Fix Session', 2)
    new_doc.add_paragraph(
        'Use this template when troubleshooting issues or fixing bugs.'
    )
    new_doc.add_paragraph()
    new_doc.add_paragraph('[Paste AIXORD Governance File]')
    new_doc.add_paragraph('PMERIT CONTINUE')
    new_doc.add_paragraph()
    new_doc.add_paragraph('My project objective: Fix [specific issue description]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('Here is the error/problem:')
    new_doc.add_paragraph('[Paste error message or describe symptoms]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('Here is the relevant code/context:')
    new_doc.add_paragraph('[Paste relevant information]')
    new_doc.add_paragraph()
    new_doc.add_paragraph('[Review AI analysis]')
    new_doc.add_paragraph('APPROVED')
    new_doc.add_paragraph()
    new_doc.add_paragraph('[Implement fix]')
    new_doc.add_paragraph('VERIFY: fix resolves the original issue')
    new_doc.add_paragraph()

    new_doc.add_page_break()

    # =========================================================================
    # APPENDIX D: Platform-Specific Tips (NEW - pages 32-33)
    # =========================================================================
    new_doc.add_heading('Appendix D: Platform-Specific Tips', 1)

    new_doc.add_paragraph(
        'While AIXORD works with any AI, each platform has unique characteristics. '
        'These tips help you get the most from each one.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Claude (Anthropic)', 2)
    new_doc.add_paragraph(
        'Claude excels at nuanced understanding and following complex instructions. '
        'It responds well to the governance framework and maintains consistency throughout sessions. '
        'Use Claude when you need careful analysis, writing quality, or code explanation. '
        'Claude handles long contexts well, making it ideal for multi-SCOPE projects. '
        'The Claude Pro tier provides extended conversation limits useful for complex sessions.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('ChatGPT (OpenAI)', 2)
    new_doc.add_paragraph(
        'ChatGPT is widely accessible and handles diverse tasks well. '
        'The GPT-4 version follows AIXORD governance more reliably than GPT-3.5. '
        'Use ChatGPT Plus for the best experience with AIXORD. '
        'ChatGPT sometimes needs reminders about behavioral firewalls, especially default suppression. '
        'The PROTOCOL CHECK command is particularly useful with ChatGPT when you notice drift.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Gemini (Google)', 2)
    new_doc.add_paragraph(
        'Gemini integrates well with Google Workspace products. '
        'It excels at tasks involving Google Docs, Sheets, and other Google services. '
        'Gemini handles factual queries and research tasks effectively. '
        'For code-heavy projects, combine Gemini with Google Colab for execution. '
        'Gemini Advanced provides the context length needed for full governance documents.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('DeepSeek', 2)
    new_doc.add_paragraph(
        'DeepSeek provides strong reasoning capabilities at competitive pricing. '
        'It handles mathematical and logical problems particularly well. '
        'DeepSeek Coder variant is optimized for programming tasks. '
        'Use the deep thinking mode for complex planning phases. '
        'DeepSeek works well for technical projects requiring step-by-step reasoning traces.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Microsoft Copilot', 2)
    new_doc.add_paragraph(
        'Copilot integrates with Microsoft 365 applications natively. '
        'Use it for tasks involving Word, Excel, PowerPoint, and Outlook. '
        'The enterprise version maintains conversation history across sessions. '
        'Copilot access to Bing search provides up-to-date information retrieval. '
        'For best results, use Copilot Pro with GPT-4 access enabled.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Cross-Platform Projects', 2)
    new_doc.add_paragraph(
        'AIXORD handoffs transfer seamlessly between platforms. You might start a project '
        'with Claude for planning, switch to ChatGPT for content generation, and finish '
        'with DeepSeek for code implementation. The handoff document preserves context '
        'regardless of which AI reads it. This flexibility lets you leverage each platform\'s '
        'strengths while maintaining project continuity.'
    )
    new_doc.add_paragraph()

    new_doc.add_heading('Universal Compatibility', 2)
    new_doc.add_paragraph(
        'Any AI that can read and follow instructions can use AIXORD governance. '
        'New AI models released in the future will work with the framework. '
        'The governance document is self-contained and platform-agnostic. '
        'If an AI can understand natural language, it can operate under AIXORD rules. '
        'This future-proofs your investment in learning the framework.'
    )

    new_doc.add_page_break()

    # =========================================================================
    # ADDITIONAL NOTES PAGE
    # =========================================================================
    new_doc.add_heading('Session Log', 1)

    new_doc.add_paragraph(
        'Track your AIXORD sessions here. Record dates, platforms, and key outcomes.'
    )
    new_doc.add_paragraph()

    for i in range(1, 11):
        new_doc.add_paragraph(f'Session {i}')
        new_doc.add_paragraph('Date: _______ Platform: _______ Project: _______')
        new_doc.add_paragraph('Key Outcome: ' + '_' * 40)
        new_doc.add_paragraph()

    new_doc.add_page_break()

    # =========================================================================
    # FINAL PAGE
    # =========================================================================
    new_doc.add_paragraph()
    new_doc.add_paragraph()
    new_doc.add_paragraph()

    final = new_doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.add_run('AIXORD v3.2.1').bold = True

    new_doc.add_paragraph()

    p = new_doc.add_paragraph('Purpose-Bound. Disciplined. Focused.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_paragraph()

    p = new_doc.add_paragraph('Copyright 2026 PMERIT LLC. All Rights Reserved.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_paragraph()
    new_doc.add_paragraph()

    p = new_doc.add_paragraph('For companion templates visit:')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = new_doc.add_paragraph('pmerit.gumroad.com')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    new_doc.add_paragraph()

    p = new_doc.add_paragraph('Use discount code: AX-STR-7K9M')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    # Save the document
    new_doc.save('AIXORD_STARTER_GUIDE_V2.docx')
    print("Expanded Starter Guide saved as AIXORD_STARTER_GUIDE_V2.docx")

    # Count words
    word_count = 0
    for para in new_doc.paragraphs:
        word_count += len(para.text.split())

    estimated_pages = word_count / 275
    print(f"Word count: {word_count}")
    print(f"Estimated pages (6x9 format, ~275 words/page): {estimated_pages:.1f}")


if __name__ == "__main__":
    expand_starter_guide()
