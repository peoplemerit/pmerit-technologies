"""
Convert Markdown manuscripts from staging/ to KDP-formatted DOCX in output/

Usage:
    python convert_md_to_kdp.py                    # Convert all manuscripts
    python convert_md_to_kdp.py AIXORD_FOR_CLAUDE  # Convert specific manuscript
"""

import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# Paths
SCRIPT_DIR = Path(__file__).parent
STAGING_DIR = SCRIPT_DIR.parent / "staging"
OUTPUT_DIR = SCRIPT_DIR.parent / "output"


def setup_kdp_styles(doc):
    """Configure document styles for KDP 6x9 format"""
    # Set page size to 6x9 inches
    for section in doc.sections:
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    return doc


def parse_markdown(md_content):
    """Parse markdown into structured sections"""
    lines = md_content.split('\n')
    sections = []
    current_section = {'type': 'text', 'content': [], 'level': 0}

    for line in lines:
        # Heading detection
        if line.startswith('# '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {'type': 'h1', 'content': line[2:].strip(), 'level': 1}
            sections.append(current_section)
            current_section = {'type': 'text', 'content': [], 'level': 0}
        elif line.startswith('## '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {'type': 'h2', 'content': line[3:].strip(), 'level': 2}
            sections.append(current_section)
            current_section = {'type': 'text', 'content': [], 'level': 0}
        elif line.startswith('### '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {'type': 'h3', 'content': line[4:].strip(), 'level': 3}
            sections.append(current_section)
            current_section = {'type': 'text', 'content': [], 'level': 0}
        elif line.startswith('#### '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {'type': 'h4', 'content': line[5:].strip(), 'level': 4}
            sections.append(current_section)
            current_section = {'type': 'text', 'content': [], 'level': 0}
        elif line.startswith('---'):
            if current_section['content']:
                sections.append(current_section)
            sections.append({'type': 'hr', 'content': '', 'level': 0})
            current_section = {'type': 'text', 'content': [], 'level': 0}
        elif line.startswith('- ') or line.startswith('* '):
            current_section['content'].append(('bullet', line[2:].strip()))
        elif re.match(r'^\d+\. ', line):
            match = re.match(r'^(\d+)\. (.+)', line)
            if match:
                current_section['content'].append(('number', match.group(2).strip()))
        elif line.startswith('**') and line.endswith('**'):
            current_section['content'].append(('bold_line', line[2:-2].strip()))
        elif line.strip() == '':
            if current_section['content'] and current_section['content'][-1] != ('blank',):
                current_section['content'].append(('blank',))
        else:
            current_section['content'].append(('text', line))

    if current_section['content']:
        sections.append(current_section)

    return sections


def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic)"""
    # Pattern for **bold** and *italic*
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def convert_to_docx(md_content, output_path, title):
    """Convert markdown content to DOCX"""
    doc = Document()
    setup_kdp_styles(doc)

    sections = parse_markdown(md_content)

    for section in sections:
        if section['type'] == 'h1':
            heading = doc.add_heading(section['content'], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif section['type'] == 'h2':
            doc.add_heading(section['content'], level=1)
        elif section['type'] == 'h3':
            doc.add_heading(section['content'], level=2)
        elif section['type'] == 'h4':
            doc.add_heading(section['content'], level=3)
        elif section['type'] == 'hr':
            doc.add_paragraph('─' * 40)
        elif section['type'] == 'text':
            paragraph_text = []
            for item in section['content']:
                if item == ('blank',):
                    if paragraph_text:
                        p = doc.add_paragraph()
                        add_formatted_text(p, ' '.join(paragraph_text))
                        paragraph_text = []
                    doc.add_paragraph()
                elif item[0] == 'bullet':
                    if paragraph_text:
                        p = doc.add_paragraph()
                        add_formatted_text(p, ' '.join(paragraph_text))
                        paragraph_text = []
                    p = doc.add_paragraph(style='List Bullet')
                    add_formatted_text(p, item[1])
                elif item[0] == 'number':
                    if paragraph_text:
                        p = doc.add_paragraph()
                        add_formatted_text(p, ' '.join(paragraph_text))
                        paragraph_text = []
                    p = doc.add_paragraph(style='List Number')
                    add_formatted_text(p, item[1])
                elif item[0] == 'bold_line':
                    if paragraph_text:
                        p = doc.add_paragraph()
                        add_formatted_text(p, ' '.join(paragraph_text))
                        paragraph_text = []
                    p = doc.add_paragraph()
                    run = p.add_run(item[1])
                    run.bold = True
                elif item[0] == 'text':
                    paragraph_text.append(item[1])

            if paragraph_text:
                p = doc.add_paragraph()
                add_formatted_text(p, ' '.join(paragraph_text))

    # Save
    doc.save(output_path)

    # Count words
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    page_estimate = word_count / 275  # ~275 words per 6x9 page

    return word_count, page_estimate


def get_staging_files():
    """Get list of markdown files in staging directory"""
    if not STAGING_DIR.exists():
        print(f"ERROR: Staging directory not found: {STAGING_DIR}")
        return []

    return list(STAGING_DIR.glob("*.md"))


def main():
    # Ensure output directory exists
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Get files to convert
    staging_files = get_staging_files()

    if not staging_files:
        print("No markdown files found in staging directory")
        return

    # Filter if specific file requested
    if len(sys.argv) > 1:
        search_term = sys.argv[1].lower()
        staging_files = [f for f in staging_files if search_term in f.stem.lower()]
        if not staging_files:
            print(f"No files matching '{sys.argv[1]}' found")
            return

    print(f"Converting {len(staging_files)} manuscript(s) to KDP format...")
    print("=" * 60)

    results = []
    for md_file in staging_files:
        # Read markdown
        md_content = md_file.read_text(encoding='utf-8')

        # Generate output filename
        output_name = md_file.stem.replace('_Manuscript', '').replace('_v4.2', '') + '_KDP.docx'
        output_path = OUTPUT_DIR / output_name

        # Convert
        try:
            word_count, page_estimate = convert_to_docx(md_content, str(output_path), md_file.stem)
            status = "OK" if page_estimate >= 24 else f"NEEDS {24 - page_estimate:.0f}+ PAGES"
            print(f"  {output_name}: {word_count:,} words, ~{page_estimate:.1f} pages - {status}")
            results.append((output_name, word_count, page_estimate, True))
        except Exception as e:
            print(f"  {output_name}: ERROR - {e}")
            results.append((output_name, 0, 0, False))

    print("=" * 60)

    # Summary
    successful = sum(1 for r in results if r[3])
    print(f"\nConverted: {successful}/{len(results)} files")
    print(f"Output directory: {OUTPUT_DIR}")

    # Check for files needing more content
    short_files = [r for r in results if r[3] and r[2] < 24]
    if short_files:
        print(f"\nWARNING: {len(short_files)} file(s) under 24 pages (KDP minimum)")


if __name__ == '__main__':
    main()
