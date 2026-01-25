"""
Expand all AIXORD manuscripts to meet 24-page KDP minimum
Target: 6,600+ words for 24 pages at 6x9 format
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Common content sections that apply to all manuscripts
def add_front_matter(doc, title, subtitle, platform_name=None):
    """Add title page, copyright, dedication"""
    # Title Page
    for _ in range(3):
        doc.add_paragraph()

    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    st = doc.add_paragraph(subtitle)
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2):
        doc.add_paragraph()

    tagline = doc.add_paragraph('Transform Chaotic AI Conversations')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tagline2 = doc.add_paragraph('Into Structured, Productive Projects')
    tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3):
        doc.add_paragraph()

    version = doc.add_paragraph('Version 3.2.1')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    publisher = doc.add_paragraph('PMERIT LLC')
    publisher.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Copyright Page
    for _ in range(2):
        doc.add_paragraph()

    doc.add_paragraph(title)
    doc.add_paragraph('Version 3.2.1')
    doc.add_paragraph()
    doc.add_paragraph('Copyright 2026 PMERIT LLC')
    doc.add_paragraph('All Rights Reserved.')
    doc.add_paragraph()
    doc.add_paragraph(
        'No part of this publication may be reproduced, distributed, or transmitted '
        'in any form or by any means, including photocopying, recording, or other '
        'electronic or mechanical methods, without the prior written permission of '
        'the publisher, except in the case of brief quotations embodied in critical '
        'reviews and certain other noncommercial uses permitted by copyright law.'
    )
    doc.add_paragraph()
    doc.add_paragraph('For permission requests, contact:')
    doc.add_paragraph('support@pmerit.com')
    doc.add_paragraph()
    doc.add_paragraph('Published by PMERIT LLC')
    doc.add_paragraph('First Edition: January 2026')

    doc.add_page_break()

    # Dedication
    for _ in range(4):
        doc.add_paragraph()

    ded = doc.add_paragraph()
    ded.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ded.add_run('For everyone who has ever felt like they were fighting their AI').italic = True

    ded2 = doc.add_paragraph()
    ded2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ded2.add_run('instead of collaborating with it.').italic = True

    for _ in range(2):
        doc.add_paragraph()

    ded3 = doc.add_paragraph()
    ded3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ded3.add_run('This framework exists because productive AI collaboration').italic = True

    ded4 = doc.add_paragraph()
    ded4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ded4.add_run("shouldn't require a PhD in prompt engineering.").italic = True

    doc.add_page_break()


def add_preface(doc, platform_name=None):
    """Add preface section"""
    doc.add_heading('Preface', 1)

    doc.add_paragraph(
        'When I first started using AI assistants for real work, I was excited by the possibilities. '
        'Here was a tool that could write code, analyze data, draft documents, and answer complex questions. '
        'The demos were impressive. The potential seemed unlimited.'
    )
    doc.add_paragraph()

    doc.add_paragraph(
        'Then reality set in. My conversations with AI became frustrating exercises in repetition. '
        'I would explain my project goals, only to have them forgotten twenty messages later. '
        'I would ask for one solution, and receive five options I did not want. '
        'I would provide specific instructions, only to watch the AI do something else entirely. '
        'The more I used these tools, the more I felt like I was working against them rather than with them.'
    )
    doc.add_paragraph()

    doc.add_paragraph(
        'The breakthrough came when I realized the problem was not the AI itself. Modern language models '
        'are remarkably capable. They can follow complex instructions when those instructions are structured correctly. '
        'What was missing was a governance framework - a set of rules that defined roles, preserved context, '
        'and controlled behavior. Without this structure, AI conversations drift into chaos.'
    )
    doc.add_paragraph()

    doc.add_paragraph(
        'AIXORD is the result of months of experimentation, iteration, and refinement. It represents '
        'everything I learned about making AI assistants genuinely useful. Every rule in the framework '
        'exists to solve a real problem I encountered. Every protocol addresses a frustration I experienced.'
    )
    doc.add_paragraph()

    if platform_name:
        doc.add_paragraph(
            f'This guide is specifically optimized for {platform_name} users. While AIXORD works with any AI platform, '
            f'each has unique characteristics that affect how governance should be applied. This book covers both the '
            f'universal AIXORD principles and the specific techniques that work best with {platform_name}.'
        )
        doc.add_paragraph()

    doc.add_paragraph(
        'This book distills the essential knowledge you need to begin using AIXORD immediately. '
        'More comprehensive documentation exists for those who want to explore advanced features, but '
        'this book contains everything required to transform your AI interactions from frustrating to productive.'
    )
    doc.add_paragraph()

    doc.add_paragraph(
        'My hope is that this framework saves you the time and frustration I experienced. AI assistants '
        'are powerful tools, but they need structure to be truly useful. AIXORD provides that structure. '
        'Welcome to productive AI collaboration.'
    )
    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run('The PMERIT Team').italic = True

    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig2.add_run('January 2026').italic = True

    doc.add_page_break()


def add_introduction(doc, platform_name=None):
    """Add introduction chapter"""
    doc.add_heading('Introduction', 1)

    doc.add_paragraph(
        'You picked up this book because you have experienced the frustration of working with '
        'AI assistants. Perhaps you have had long conversations that went nowhere. Perhaps you '
        'have lost important context when a session ended. Perhaps the AI kept suggesting things '
        'you did not ask for, or failed to follow your instructions. You are not alone.'
    )
    doc.add_paragraph()

    doc.add_paragraph(
        'The problem is not the AI itself. Modern language models are remarkably capable. '
        'They can write code, analyze data, create content, and solve complex problems. '
        'The problem is the lack of structure in how we interact with them. Without clear '
        'rules and roles, AI conversations devolve into chaos.'
    )
    doc.add_paragraph()

    doc.add_heading('What This Book Will Teach You', 2)
    doc.add_paragraph(
        'AIXORD (AI Execution Order) is a governance framework that brings order to AI collaboration. '
        'By the time you finish this book, you will know how to:'
    )
    doc.add_paragraph()

    skills = [
        'Establish clear roles where YOU make decisions and AI executes them',
        'Use phases to structure any project from idea to completion',
        'Preserve context across sessions so nothing is ever forgotten',
        'Control AI behavior with behavioral firewalls that suppress unwanted defaults',
        'Decompose large projects into manageable, verifiable pieces',
        'Verify that what was built matches what was planned'
    ]
    for skill in skills:
        doc.add_paragraph('  * ' + skill)
    doc.add_paragraph()

    doc.add_heading('Who This Book Is For', 2)
    doc.add_paragraph(
        'This book is for anyone who uses AI assistants for productive work. Whether you are '
        'a developer building software, a marketer creating campaigns, a student researching '
        'papers, or a hobbyist building personal projects, AIXORD will make your AI interactions '
        'more efficient and less frustrating.'
    )
    doc.add_paragraph()

    if platform_name:
        doc.add_paragraph(
            f'This edition is specifically written for {platform_name} users, with examples and '
            f'techniques tailored to that platform. However, the core AIXORD principles apply to '
            f'any AI system, so you can transfer your knowledge if you switch platforms.'
        )
        doc.add_paragraph()

    doc.add_page_break()


def add_what_is_aixord(doc):
    """Add 'What is AIXORD' chapter"""
    doc.add_heading('Chapter 1: What is AIXORD?', 1)

    doc.add_paragraph(
        'AIXORD (AI Execution Order) is a governance framework for human-AI collaboration. '
        'It transforms chaotic AI conversations into structured, productive project execution. '
        'Instead of hoping your AI assistant understands what you want, AIXORD ensures '
        'clear communication, documented decisions, and verified outcomes.'
    )
    doc.add_paragraph()

    doc.add_heading('The Problem with Unstructured AI Use', 2)
    doc.add_paragraph(
        'Without governance, AI conversations often become frustrating experiences. '
        'The AI forgets what you discussed earlier in the session. It offers unwanted '
        'suggestions and alternatives when you just want one answer. It goes off-topic '
        'to discuss interesting but irrelevant tangents. It makes assumptions about '
        "what you want without asking. It claims tasks are complete without verification."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'These issues compound over multi-session projects. Context gets lost between '
        'conversations. Decisions made yesterday are forgotten today. You spend more time '
        're-explaining your project than actually making progress. The AI becomes a source '
        'of frustration rather than productivity.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Consider a common scenario: You spend an hour carefully explaining your project '
        'requirements to an AI assistant. The conversation goes well. The AI seems to understand. '
        'Then you close the browser and come back the next day. The AI has no memory of your '
        'previous conversation. All that context is gone. You must start from scratch.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Or another scenario: You ask the AI for a simple recommendation. Instead of giving '
        'you one answer, it provides five options with lengthy explanations of each. You did '
        'not want options. You wanted a decision. But the AI, trying to be helpful, gave you '
        'more than you asked for and wasted your time reading through alternatives you never '
        'requested.'
    )
    doc.add_paragraph()

    doc.add_heading('The AIXORD Solution', 2)
    doc.add_paragraph(
        'AIXORD provides structure through clear roles, defined phases, and explicit rules. '
        'You become the Director who makes all decisions. The AI becomes your Architect '
        'who plans and your Commander who executes. Every decision is documented in a '
        'decision ledger. Every action requires authorization. Context persists through '
        'handoff documents that transfer complete state between sessions.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'The framework includes behavioral firewalls that suppress unwanted AI defaults. '
        'No more verbose explanations when you want concise answers. No more multiple '
        'options when you asked for one recommendation. No more "anything else?" prompts '
        'when the task is complete.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'With AIXORD, the AI operates under explicit governance. It knows its role. It knows '
        'the boundaries. It knows what it can and cannot do without your permission. This '
        'structure transforms the AI from an unpredictable conversation partner into a '
        'reliable project execution tool.'
    )
    doc.add_paragraph()

    doc.add_heading('Core Principle', 2)
    doc.add_paragraph(
        'The fundamental principle of AIXORD is simple: You (Human) are the Director. '
        'AI is your Architect and Commander. Every decision is documented, every action '
        'is authorized, and nothing is forgotten between sessions. This principle guides '
        'every feature and rule in the framework.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'This is not about limiting AI capabilities. Modern AI systems are remarkably capable. '
        'The problem is not capability but direction. AIXORD provides that direction. It channels '
        'AI capability toward your specific goals rather than letting it wander in whatever '
        'direction seems interesting to the model.'
    )
    doc.add_paragraph()

    doc.add_heading('The AIXORD Formula', 2)
    doc.add_paragraph(
        'AIXORD uses a project composition formula to structure work:'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Steps (Seconds) -> Deliverables (Minutes) -> Master Scope (The Hour) = Done').italic = True
    doc.add_paragraph()
    doc.add_paragraph(
        'Small actions build deliverables. Deliverables build the complete project. '
        'This decomposition approach makes even large projects manageable by breaking '
        'them into trackable, verifiable pieces called SCOPEs. Each SCOPE has clear '
        'requirements, clear completion criteria, and clear verification methods.'
    )

    doc.add_page_break()


def add_quick_start(doc, platform_name=None, platform_url=None):
    """Add Quick Start chapter"""
    doc.add_heading('Chapter 2: Quick Start Guide', 1)

    doc.add_paragraph(
        'Get productive with AIXORD in 15 minutes. This chapter walks you through '
        'your first AIXORD session from start to finish.'
    )
    doc.add_paragraph()

    doc.add_heading('Before You Begin', 2)
    doc.add_paragraph('Make sure you have the following ready:')
    prereqs = [
        'The AIXORD governance file (included with your purchase)',
        f'Access to {platform_name or "your AI"} ({platform_url or "your AI chat interface"})',
        'A project idea or task you want to accomplish',
        'Your purchase email for license validation'
    ]
    for p in prereqs:
        doc.add_paragraph('  * ' + p)
    doc.add_paragraph()

    doc.add_heading('Step 1: Copy the Governance File', 2)
    doc.add_paragraph(
        'Open the AIXORD governance file included with your purchase. '
        'This file contains all the rules that will govern your AI session. Select all the text '
        '(Ctrl+A on Windows, Cmd+A on Mac) and copy it (Ctrl+C or Cmd+C).'
    )
    doc.add_paragraph()

    doc.add_heading('Step 2: Paste into Your AI', 2)
    doc.add_paragraph(
        f'Open {platform_name or "your preferred AI chat interface"}. '
        'Start a new conversation (do not continue an existing one). Paste the governance '
        'document into the input field and press Enter to send.'
    )
    doc.add_paragraph(
        'The AI will read and acknowledge the governance framework. It will display a '
        'confirmation message and begin the setup process.'
    )
    doc.add_paragraph()

    doc.add_heading('Step 3: Complete License Validation', 2)
    doc.add_paragraph(
        'The AI will ask for your license email or authorization code. Enter the email '
        'address you used when purchasing AIXORD. The system validates your license and '
        'confirms you are authorized to use the framework.'
    )
    doc.add_paragraph()

    doc.add_heading('Step 4: Answer Setup Questions', 2)
    doc.add_paragraph('The AI will ask several questions to configure your session:')
    setup_questions = [
        'Which AI platform are you using?',
        'What tier is your account? (Free, Plus/Pro, or Enterprise)',
        'What is your project objective? (1-2 sentences describing what you want to build)'
    ]
    for q in setup_questions:
        doc.add_paragraph('  * ' + q)
    doc.add_paragraph()

    doc.add_heading('Step 5: Start Working', 2)
    doc.add_paragraph(
        'Once setup is complete, you are ready to work. The AI will display a response header '
        'showing your current phase, progress, and project scope.'
    )

    doc.add_page_break()


def add_authority_model(doc):
    """Add Authority Model chapter"""
    doc.add_heading('Chapter 3: The Authority Model', 1)

    doc.add_paragraph(
        'AIXORD establishes clear roles to prevent confusion about who decides what. '
        'This clarity is essential for productive collaboration. Without defined roles, '
        'AI and human work at cross-purposes, leading to wasted effort and frustration.'
    )
    doc.add_paragraph()

    doc.add_heading('The Director (You)', 2)
    doc.add_paragraph(
        'You are the Director. This is the most important role in AIXORD. As Director, '
        'you make ALL decisions. You approve ALL execution. You own ALL outcomes. '
        'The AI cannot proceed with any implementation without your explicit approval.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'The Director role means you are in control. You decide what to build, how to build it, '
        'and when work is complete. The AI serves you, not the other way around. If you '
        'disagree with an AI recommendation, your decision stands. If you want to change '
        'direction mid-project, you change direction. The framework protects your authority.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'As Director, you set the project objective. You approve or reject proposals. You decide '
        'when to move between phases. You determine when work is complete. Every significant '
        'action requires your explicit authorization. The AI waits for your direction rather '
        'than taking initiative on its own.'
    )
    doc.add_paragraph()

    doc.add_heading('The Architect (AI in Planning Mode)', 2)
    doc.add_paragraph(
        'When planning and specifying, the AI serves as your Architect. In this role, '
        'the AI analyzes problems, asks clarifying questions, researches options, and '
        'creates specifications. The Architect recommends but never decides.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'As Architect, the AI produces project documents, scope specifications, and '
        'implementation plans. It identifies risks and suggests mitigations. It breaks '
        'large projects into manageable deliverables. All of this work requires your '
        'approval before any implementation begins.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'The Architect role is about preparation and planning. The AI thinks through '
        'problems carefully, considers alternatives, and presents you with well-reasoned '
        'recommendations. But the final decision always rests with you as Director.'
    )
    doc.add_paragraph()

    doc.add_heading('The Commander (AI in Execution Mode)', 2)
    doc.add_paragraph(
        'When you say APPROVED, the AI transitions to Commander role. The Commander '
        'implements your approved plans. It writes code, creates content, and delivers '
        'artifacts. It reports progress and surfaces blockers for your decision.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'The Commander stays within the approved scope. It does not expand requirements '
        "or add features without your approval. If it encounters something unexpected, "
        "it stops and asks for direction. Implementation follows specification exactly."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'During execution, the Commander provides regular progress updates. It tells you '
        'what has been completed, what is in progress, and what remains. If it encounters '
        'a blocker, it immediately surfaces the issue for your decision rather than trying '
        'to work around it independently.'
    )
    doc.add_paragraph()

    doc.add_heading('The Golden Rule', 2)
    doc.add_paragraph(
        'Authority flows in one direction: Decisions flow DOWN from Director to Architect '
        'to Commander. Implementation flows UP from Commander to Architect to Director '
        'for approval. This prevents the AI from acting unilaterally and ensures you '
        'remain in control of your project at all times.'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Director -> Architect -> Commander (decisions down)').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Commander -> Architect -> Director (implementation up)').bold = True
    doc.add_paragraph()
    doc.add_paragraph(
        'This bidirectional flow ensures that decisions are made at the appropriate level '
        'and that implementation is always verified against the original intent. Nothing '
        'falls through the cracks because every action has a clear owner and approval chain.'
    )

    doc.add_page_break()


def add_phase_system(doc):
    """Add Phase System chapter"""
    doc.add_heading('Chapter 4: The Phase System', 1)

    doc.add_paragraph(
        'AIXORD uses six distinct phases to structure your workflow. Each phase has '
        'specific rules about what the AI can and cannot do. This prevents the AI from '
        'jumping ahead or working on the wrong thing at the wrong time.'
    )
    doc.add_paragraph()

    doc.add_heading('DECISION Phase', 2)
    doc.add_paragraph(
        'The default waiting state. The AI awaits your direction and does not take '
        'initiative. This is the safe harbor where you can think, plan, and decide '
        'without the AI rushing you toward execution. When you need time to consider '
        'your options, the AI patiently waits in DECISION phase.'
    )
    doc.add_paragraph()

    doc.add_heading('DISCOVER Phase', 2)
    doc.add_paragraph(
        "For unclear ideas. When you do not know exactly what you want to build, "
        'DISCOVER mode activates exploratory questioning. The AI asks probing questions '
        "to help you clarify your vision. It does not suggest solutions yet - it helps "
        'you understand the problem first. This is particularly valuable when starting '
        'a new project or exploring unfamiliar territory.'
    )
    doc.add_paragraph()

    doc.add_heading('BRAINSTORM Phase', 2)
    doc.add_paragraph(
        'For generating possibilities. Once you have a clearer problem statement, '
        'BRAINSTORM mode enables creative exploration. The AI generates ideas, '
        'considers approaches, and helps you think through options. Output is '
        'intentionally expansive at this stage. Judgment is suspended in favor of '
        'creativity. The goal is to explore the solution space broadly before narrowing.'
    )
    doc.add_paragraph()

    doc.add_heading('OPTIONS Phase', 2)
    doc.add_paragraph(
        'For comparing approaches. When ready to choose a path, OPTIONS mode presents '
        '2-3 clear alternatives with pros, cons, costs, and timelines. The AI helps '
        'you make an informed decision by structuring the comparison. Each option is '
        'presented with enough detail to understand implications but not so much detail '
        'that the comparison becomes overwhelming.'
    )
    doc.add_paragraph()

    doc.add_heading('EXECUTE Phase', 2)
    doc.add_paragraph(
        'For implementation. This phase requires explicit APPROVED command to enter. '
        "The AI implements your approved plan step by step. It does not improvise or "
        'expand scope. It follows the specification you approved. Progress is tracked '
        'and reported. Blockers are surfaced immediately for your decision.'
    )
    doc.add_paragraph()

    doc.add_heading('AUDIT Phase', 2)
    doc.add_paragraph(
        'For reviewing work. After implementation, AUDIT mode activates verification. '
        'The AI compares completed work against specifications and identifies '
        'discrepancies. This ensures what was built matches what was planned. No '
        'implementation is considered complete until it passes audit.'
    )
    doc.add_paragraph()

    doc.add_heading('Critical Rule: EXECUTE Requires APPROVED', 2)
    doc.add_paragraph(
        'The most important phase rule: EXECUTE requires explicit "APPROVED" command. '
        'The AI never implements without your permission. This prevents the AI from '
        'building the wrong thing. It also gives you a moment to review the plan '
        'before committing to implementation.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "If you are not ready to approve, say so. Ask more questions. Request changes. "
        "The AI will remain in planning mode until you explicitly approve. There is no "
        "pressure to move forward until you are confident in the plan."
    )

    doc.add_page_break()


def add_session_continuity(doc):
    """Add Session Continuity chapter"""
    doc.add_heading('Chapter 5: Session Continuity', 1)

    doc.add_paragraph(
        'AIXORD ensures your work persists between sessions. This is one of the most '
        'valuable features of the framework, especially for multi-day projects that '
        'span many conversations.'
    )
    doc.add_paragraph()

    doc.add_heading('The Problem with AI Memory', 2)
    doc.add_paragraph(
        'Most AI systems have no memory between conversations. When you close a chat '
        'and open a new one, the AI has no idea what you discussed before. All your '
        'context, decisions, and progress are lost. You start from zero every time.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'This creates enormous inefficiency for real projects. You spend the first '
        'ten minutes of every session re-explaining what you already explained yesterday. '
        'Decisions get made and then forgotten. Work gets duplicated because the AI does '
        'not remember what was already completed. Progress feels impossible.'
    )
    doc.add_paragraph()

    doc.add_heading('The HANDOFF Protocol', 2)
    doc.add_paragraph(
        'Before ending a session, type PMERIT HANDOFF. The AI generates a complete summary '
        'of your session including: current phase and status, all decisions made during '
        'the session, work completed, work remaining, blockers encountered, and any '
        'questions that need answers.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Copy the handoff document and save it as HANDOFF_[DATE].md in your project folder. '
        "This becomes your session's permanent record. Even if the AI forgets, you have "
        'documentation of everything that happened. This is your insurance against lost context.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'The handoff document is structured for machine readability. When you paste it into '
        'a new session, the AI can parse it and reconstruct the complete context. It knows '
        'what phase you were in, what decisions were made, and what comes next.'
    )
    doc.add_paragraph()

    doc.add_heading('Resuming Work', 2)
    doc.add_paragraph(
        'Next session, follow this sequence: First, paste the AIXORD governance file. '
        'Second, paste your most recent handoff document. Third, type PMERIT CONTINUE. '
        'The AI reads both documents and reconstructs the complete context.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'This three-step sequence is essential. The governance file establishes the rules. '
        'The handoff provides the context. The CONTINUE command activates the framework. '
        'Skip any step and the session will not work correctly.'
    )
    doc.add_paragraph()

    doc.add_heading('CHECKPOINT for Mid-Session Saves', 2)
    doc.add_paragraph(
        'Use CHECKPOINT to save state without ending your session. This is useful for long '
        'sessions when you want to preserve progress but are not ready to stop. The AI '
        'generates a checkpoint document you can save as a safety backup.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Checkpoints are especially valuable when you are about to try something risky '
        'or experimental. Save a checkpoint first, then proceed. If things go wrong, '
        'you can restore from the checkpoint rather than losing all your progress.'
    )
    doc.add_paragraph()

    doc.add_heading('Session Duration Recommendations', 2)
    doc.add_paragraph(
        'For best results, keep sessions focused and time-bounded. A session of 30-90 '
        'minutes is ideal. Longer sessions risk context degradation as the conversation '
        'grows. Shorter sessions may not provide enough time for meaningful progress.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'When you see warnings about context length or when responses start becoming '
        'less coherent, it is time to HANDOFF and start fresh. Do not push past these '
        'limits or you risk losing important context.'
    )

    doc.add_page_break()


def add_purpose_bound(doc):
    """Add Purpose-Bound Operation chapter"""
    doc.add_heading('Chapter 6: Purpose-Bound Operation', 1)

    doc.add_paragraph(
        'Once you state your project objective, the AI commits exclusively to that scope. '
        'This prevents session drift and keeps both you and the AI focused on what matters.'
    )
    doc.add_paragraph()

    doc.add_heading('The Commitment', 2)
    p = doc.add_paragraph()
    p.add_run(
        '"I exist in this session ONLY to serve your stated project objective. '
        'I will not engage with topics outside that scope unless you explicitly expand it."'
    ).italic = True
    doc.add_paragraph()
    doc.add_paragraph(
        'This commitment transforms how the AI operates. It will not answer random questions '
        'unrelated to your project. It will not go on tangents about interesting but '
        'irrelevant topics. It stays focused on helping you complete your stated objective.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Purpose-bound operation is like having a dedicated assistant who only cares about '
        'your project. Everything else is noise. This focus dramatically improves productivity '
        'because neither you nor the AI waste time on distractions.'
    )
    doc.add_paragraph()

    doc.add_heading('Enforcement Levels', 2)
    doc.add_paragraph(
        'Purpose-bound operation has three enforcement levels. Choose the level that matches '
        'your working style and project requirements.'
    )
    doc.add_paragraph()

    doc.add_heading('STRICT Enforcement', 3)
    doc.add_paragraph(
        'Zero tolerance for off-topic requests. Immediate redirect without acknowledgment. '
        'Use STRICT when you need maximum focus, when working on deadline, or when the '
        'project is complex enough that any distraction is costly.'
    )
    doc.add_paragraph()

    doc.add_heading('STANDARD Enforcement', 3)
    doc.add_paragraph(
        'Default behavior. Polite redirect explaining why the request is outside scope, '
        'with options to expand scope or return to project. STANDARD works well for most '
        'situations where you want focus but also flexibility.'
    )
    doc.add_paragraph()

    doc.add_heading('RELAXED Enforcement', 3)
    doc.add_paragraph(
        'Brief tangent allowed for context, then redirect back to project. Use RELAXED '
        'when exploring related topics might provide useful context, or when the project '
        'boundaries are still being defined.'
    )
    doc.add_paragraph()

    doc.add_heading('Expanding Scope', 2)
    doc.add_paragraph(
        'If you need to add a new topic to your project, use EXPAND SCOPE: [topic]. '
        'The AI will acknowledge the expansion and incorporate it into the project boundaries. '
        "This is different from going off-topic - it is an intentional broadening of your "
        'project that becomes part of the documented scope.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Scope expansions are recorded in the decision ledger. This provides an audit trail '
        'of how your project evolved over time. You can see exactly when and why scope was '
        'added, which helps with project management and retrospectives.'
    )

    doc.add_page_break()


def add_behavioral_firewalls(doc):
    """Add Behavioral Firewalls chapter"""
    doc.add_heading('Chapter 7: Behavioral Firewalls', 1)

    doc.add_paragraph(
        'AIXORD controls unwanted AI behaviors through a system of behavioral firewalls. '
        'These rules suppress the AI defaults that cause frustration and waste time.'
    )
    doc.add_paragraph()

    doc.add_heading('Default Suppression', 2)
    doc.add_paragraph(
        'Unless explicitly requested, the AI suppresses: detailed explanations, '
        'examples, suggestions, alternatives, comparisons, and future considerations. '
        'You get only what you ask for. Nothing extra. Nothing unsolicited.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'This is the opposite of how most AI systems behave by default. Normally, an AI '
        'tries to be helpful by providing extra context, examples, and alternatives. But '
        'this helpfulness often becomes overwhelming. AIXORD flips the default: minimal '
        'output unless you explicitly ask for more.'
    )
    doc.add_paragraph()

    doc.add_heading('Choice Elimination', 2)
    doc.add_paragraph(
        'The AI provides ONE answer, not multiple options, unless you specifically request alternatives. '
        "No more 'Here are 5 ways to do this...' when you just want one good solution. "
        'If you want options, ask for them explicitly using the OPTIONS phase.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Choice elimination respects your time. When you ask a question, you usually want '
        'an answer, not a menu. The AI makes its best recommendation and delivers it cleanly. '
        'If you disagree, you can ask for alternatives. But the default is decisive action.'
    )
    doc.add_paragraph()

    doc.add_heading('Expansion Triggers', 2)
    doc.add_paragraph(
        'Use trigger words to permit detailed output when you want it: EXPLAIN, WHY, '
        'TEACH, DETAIL, OPTIONS, COMPARE, ELABORATE. These words signal to the AI that '
        'verbose output is acceptable. Without these words, responses stay minimal.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Trigger words put you in control of output verbosity. When you need a quick answer, '
        'you get a quick answer. When you need a deep explanation, you ask for one. The AI '
        'adapts to your needs rather than imposing its own idea of appropriate detail.'
    )
    doc.add_paragraph()

    doc.add_heading('Hard Stop', 2)
    doc.add_paragraph(
        'After completing a task, the AI stops immediately. No follow-up questions. '
        'No "anything else I can help with?" prompts. No suggestions for what to do next. '
        'Task done equals response ends. You control when the next task begins.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Hard stop eliminates the conversational filler that plagues most AI interactions. '
        'When the AI finishes a task, it is done. It does not try to keep the conversation '
        'going. It does not fish for more work. It waits quietly for your next instruction.'
    )
    doc.add_paragraph()

    doc.add_heading('Instruction Priority Stack', 2)
    doc.add_paragraph(
        'When instructions conflict, AIXORD follows a clear priority order: Governance rules '
        'take precedence over user commands. User commands take precedence over task content. '
        'Task content takes precedence over default AI training. This hierarchy ensures '
        'predictable behavior even in complex situations.'
    )

    doc.add_page_break()


def add_project_decomposition(doc):
    """Add Project Decomposition chapter"""
    doc.add_heading('Chapter 8: Project Decomposition', 1)

    doc.add_paragraph(
        'AIXORD breaks projects into manageable, trackable pieces using formal decomposition.'
    )
    doc.add_paragraph()

    doc.add_heading('The System Equation', 2)
    p = doc.add_paragraph()
    p.add_run('MASTER_SCOPE = Project_Docs = All_SCOPEs = Production-Ready System').bold = True
    doc.add_paragraph()

    doc.add_heading('Hierarchy Structure', 2)
    doc.add_paragraph('MASTER_SCOPE (Complete project)')
    doc.add_paragraph('    SCOPE_A (Deliverable 1)')
    doc.add_paragraph('        SUB-SCOPE_A1 (Step 1)')
    doc.add_paragraph('        SUB-SCOPE_A2 (Step 2)')
    doc.add_paragraph('    SCOPE_B (Deliverable 2)')
    doc.add_paragraph()

    doc.add_heading('Key Principle', 2)
    doc.add_paragraph(
        '"If it is not documented, it does not exist." You cannot implement something '
        'not specified. Documentation and reality must match.'
    )

    doc.add_page_break()


def add_locking_system(doc):
    """Add 4-State Locking System chapter"""
    doc.add_heading('Chapter 9: The 4-State Locking System', 1)

    doc.add_paragraph(
        'Every SCOPE progresses through four states, providing clear status at all times.'
    )
    doc.add_paragraph()

    states = [
        ('PLANNED', 'Plan complete, implementation not begun.'),
        ('ACTIVE', 'Under active development. Changes allowed.'),
        ('IMPLEMENTED', 'Development complete, ready for audit.'),
        ('VERIFIED', 'Audited and confirmed working. Locked.')
    ]

    for state, desc in states:
        doc.add_heading(state, 2)
        doc.add_paragraph(desc)
        doc.add_paragraph()

    doc.add_heading('Why Locking Matters', 2)
    doc.add_paragraph(
        'Locking prevents regression. Once verified, a SCOPE cannot change without '
        'explicit unlock. This protects completed work.'
    )

    doc.add_page_break()


def add_visual_audit(doc):
    """Add Visual Audit chapter"""
    doc.add_heading('Chapter 10: Visual Audit Protocol', 1)

    doc.add_paragraph(
        'AIXORD verifies implementations match specifications through formal audit.'
    )
    doc.add_paragraph()

    doc.add_heading('When Visual Audit is Required', 2)
    doc.add_paragraph('Visual audits are REQUIRED for:')
    doc.add_paragraph('  * UI Features - buttons, layouts, styling')
    doc.add_paragraph('  * Forms - input fields, validation messages')
    doc.add_paragraph('  * Dashboards - charts, data displays')
    doc.add_paragraph()

    doc.add_heading('The Audit Process', 2)
    steps = [
        ('CAPTURE', 'You provide screenshots'),
        ('COMPARE', 'AI compares against requirements'),
        ('DOCUMENT', 'Findings recorded'),
        ('VERDICT', 'PASS or DISCREPANCY')
    ]
    for step, desc in steps:
        p = doc.add_paragraph()
        p.add_run(step + ': ').bold = True
        p.add_run(desc)

    doc.add_page_break()


def add_common_mistakes(doc):
    """Add Common Mistakes chapter"""
    doc.add_heading('Chapter 11: Common Mistakes to Avoid', 1)

    doc.add_paragraph(
        'Even with AIXORD governance, users sometimes fall into patterns that reduce effectiveness. '
        "Here are the most common mistakes and how to avoid them."
    )
    doc.add_paragraph()

    doc.add_heading('Mistake 1: Skipping the Setup', 2)
    doc.add_paragraph(
        'Jumping straight into tasks without completing the setup process. This causes the AI '
        'to operate without governance context and revert to unhelpful defaults. Always complete '
        'setup. It takes 5 minutes and saves hours of frustration later.'
    )
    doc.add_paragraph()

    doc.add_heading('Mistake 2: Forgetting the Response Header', 2)
    doc.add_paragraph(
        'Not noticing when the AI stops displaying the mandatory response header. This indicates '
        'session drift without phase or progress tracking. If the header disappears for more than '
        'two responses, issue PROTOCOL CHECK immediately.'
    )
    doc.add_paragraph()

    doc.add_heading('Mistake 3: Approving Without Reviewing', 2)
    doc.add_paragraph(
        'Saying "APPROVED" before carefully reviewing the AI proposal. This results in the AI '
        'executing flawed plans that must be redone from the beginning. Read proposals fully. '
        'Ask questions. Request changes. Only then approve.'
    )
    doc.add_paragraph()

    doc.add_heading('Mistake 4: Ignoring Checkpoint Warnings', 2)
    doc.add_paragraph(
        'Pushing past message thresholds without saving session state. This leads to context '
        'degradation where the AI forgets earlier decisions and work is lost. When you see '
        '"Consider CHECKPOINT soon" do it immediately.'
    )
    doc.add_paragraph()

    doc.add_heading('Mistake 5: Scope Creep Through Quick Questions', 2)
    doc.add_paragraph(
        'Asking tangential questions that pull the AI off-project. The session becomes unfocused '
        'and productivity drops. Use PURPOSE-BOUND enforcement. Save tangents for separate sessions.'
    )
    doc.add_paragraph()

    doc.add_heading('Mistake 6: Not Using Phases', 2)
    doc.add_paragraph(
        'Staying in DECISION phase for everything without transitioning. This misses the structure '
        'that phases provide and conversations meander aimlessly. Explicitly enter DISCOVER, '
        'BRAINSTORM, OPTIONS, or EXECUTE as appropriate for your current need.'
    )
    doc.add_paragraph()

    doc.add_heading('Mistake 7: Over-Relying on AI Memory', 2)
    doc.add_paragraph(
        'Assuming the AI remembers everything from earlier in the session. This causes confusion '
        'and contradictory outputs as context degrades. Use CHECKPOINT regularly. Reference '
        'specific decisions explicitly rather than assuming the AI remembers them.'
    )
    doc.add_paragraph()

    doc.add_heading('Mistake 8: Fighting the Framework', 2)
    doc.add_paragraph(
        'Trying to "trick" or work around governance rules. This defeats the purpose of using '
        "the framework. Trust the process. If rules feel wrong, modify them - do not bypass them."
    )
    doc.add_paragraph()

    doc.add_heading('The Meta-Lesson', 2)
    doc.add_paragraph(
        'AIXORD works when you work with it. The framework is a tool, not a cage. '
        'If something is not working for your situation, adjust the framework parameters '
        'rather than abandoning governance entirely. The rules can be tuned. The principle '
        'of structured collaboration should remain.'
    )

    doc.add_page_break()


def add_commands_reference(doc):
    """Add Commands Reference chapter"""
    doc.add_heading('Chapter 12: Commands Reference', 1)

    doc.add_paragraph(
        'This chapter provides a complete reference for all AIXORD commands. Keep this '
        'chapter bookmarked for quick access during your sessions.'
    )
    doc.add_paragraph()

    doc.add_heading('Activation and Control', 2)
    doc.add_paragraph('These commands control the overall session state:')
    doc.add_paragraph()
    cmds = [
        ('PMERIT CONTINUE', 'Activate AIXORD or resume from a handoff document'),
        ('CHECKPOINT', 'Save current state without ending the session'),
        ('HANDOFF', 'Generate full session export and prepare to end'),
        ('HALT', 'Stop all current work and return to DECISION phase'),
        ('APPROVED', 'Approve current plan and enter EXECUTE phase')
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph()
        p.add_run(cmd).bold = True
        p.add_run(' - ' + desc)
    doc.add_paragraph()

    doc.add_heading('Phase Commands', 2)
    doc.add_paragraph('These commands transition between operational phases:')
    doc.add_paragraph()
    phase_cmds = [
        ('PMERIT DISCOVER', 'Enter Discovery mode for exploratory questioning'),
        ('PMERIT BRAINSTORM', 'Enter Brainstorm mode for creative ideation'),
        ('PMERIT OPTIONS', 'Request structured comparison of alternatives'),
        ('PMERIT DOCUMENT', 'Generate formal project documentation'),
        ('PMERIT EXECUTE', 'Signal readiness for implementation'),
        ('PMERIT STATUS', 'Display current phase, progress, and state')
    ]
    for cmd, desc in phase_cmds:
        p = doc.add_paragraph()
        p.add_run(cmd).bold = True
        p.add_run(' - ' + desc)
    doc.add_paragraph()

    doc.add_heading('Enforcement Commands', 2)
    doc.add_paragraph('These commands enforce governance compliance:')
    doc.add_paragraph()
    enforce_cmds = [
        ('PROTOCOL CHECK', 'Force AI to verify its own compliance with rules'),
        ('DRIFT WARNING', 'Flag when AI behavior deviates from governance'),
        ('VERIFY: [claim]', 'Request proof for a specific claim or completion'),
        ('SHOW SCOPE', 'Display current project boundaries')
    ]
    for cmd, desc in enforce_cmds:
        p = doc.add_paragraph()
        p.add_run(cmd).bold = True
        p.add_run(' - ' + desc)
    doc.add_paragraph()

    doc.add_heading('Purpose-Bound Commands', 2)
    doc.add_paragraph('These commands control focus and scope:')
    doc.add_paragraph()
    pb_cmds = [
        ('PURPOSE-BOUND: STRICT', 'Maximum enforcement, zero tolerance for tangents'),
        ('PURPOSE-BOUND: STANDARD', 'Default enforcement with redirect options'),
        ('PURPOSE-BOUND: RELAXED', 'Brief tangents allowed with redirect'),
        ('EXPAND SCOPE: [topic]', 'Add a new topic to project boundaries')
    ]
    for cmd, desc in pb_cmds:
        p = doc.add_paragraph()
        p.add_run(cmd).bold = True
        p.add_run(' - ' + desc)
    doc.add_paragraph()

    doc.add_heading('Locking Commands', 2)
    doc.add_paragraph('These commands manage SCOPE state transitions:')
    doc.add_paragraph()
    lock_cmds = [
        ('UNLOCK: [scope]', 'Transition SCOPE from PLANNED to ACTIVE'),
        ('LOCK: [scope]', 'Transition SCOPE to VERIFIED and prevent changes'),
        ('IMPLEMENTATION COMPLETE', 'Signal that SCOPE is ready for audit')
    ]
    for cmd, desc in lock_cmds:
        p = doc.add_paragraph()
        p.add_run(cmd).bold = True
        p.add_run(' - ' + desc)

    doc.add_page_break()


def add_troubleshooting(doc, platform_name=None):
    """Add Troubleshooting chapter"""
    doc.add_heading('Chapter 13: Troubleshooting FAQ', 1)

    doc.add_paragraph(
        'Even well-governed AI sessions can encounter issues. This chapter provides solutions '
        'to the most common problems AIXORD users face.'
    )
    doc.add_paragraph()

    doc.add_heading('AI provides unsolicited options', 2)
    doc.add_paragraph(
        'Issue DRIFT WARNING to remind the AI of Choice Elimination rules. If it persists, '
        'say: "PROTOCOL CHECK - Choice Elimination violated. Provide ONE recommendation, '
        'not multiple options." The AI should immediately correct its behavior and provide '
        'a single, decisive answer.'
    )
    doc.add_paragraph()

    doc.add_heading('Session context seems lost', 2)
    doc.add_paragraph(
        'Always save HANDOFF before ending any session. To resume properly, follow this '
        'exact sequence: paste governance file first, then paste your latest handoff, '
        'then type PMERIT CONTINUE. All three steps are required in that order. If you '
        'skip any step, context will not be properly restored.'
    )
    doc.add_paragraph()

    doc.add_heading('AI claims task is complete without verification', 2)
    doc.add_paragraph(
        'Issue VERIFY: [task name] to request proof of completion. Demand screenshots, '
        'test output, or other evidence. Never accept "it should work" as verification. '
        'The AI must demonstrate that the work meets specifications before any SCOPE '
        'can transition to VERIFIED status.'
    )
    doc.add_paragraph()

    doc.add_heading('AI ignores my instructions', 2)
    doc.add_paragraph(
        'Say: "PROTOCOL CHECK - Instruction Priority Stack violation. User commands override '
        'training defaults. Acknowledge and comply." This reminds the AI that your explicit '
        'instructions take precedence over its default behaviors.'
    )
    doc.add_paragraph()

    doc.add_heading('Response header keeps disappearing', 2)
    doc.add_paragraph(
        'Say: "PROTOCOL CHECK - Response header missing for [N] responses. Reinstate mandatory '
        'header immediately." The header must appear on every response to maintain visibility '
        'into session state. If it disappears repeatedly, consider starting a fresh session.'
    )
    doc.add_paragraph()

    doc.add_heading('AI keeps going off-topic', 2)
    doc.add_paragraph(
        'Issue PURPOSE-BOUND: STRICT to enforce zero tolerance for tangents. Or issue '
        'SHOW SCOPE to remind the AI of current project boundaries. For persistent issues, '
        'start a fresh session with a governance repaste. Some drift is inevitable in long '
        'sessions, which is why checkpoints and fresh starts are important.'
    )
    doc.add_paragraph()

    doc.add_heading('AI produces low-quality output', 2)
    doc.add_paragraph(
        'If output quality degrades, check whether the session has become too long. Context '
        'degradation affects output quality. Try a HANDOFF and fresh start. Also ensure you '
        'are providing clear, specific instructions. Vague requests produce vague results.'
    )
    doc.add_paragraph()

    doc.add_heading('Cannot resume from handoff', 2)
    doc.add_paragraph(
        'Verify the handoff document is complete and not truncated. Check that you pasted '
        'the governance file first. Ensure you typed PMERIT CONTINUE after pasting both '
        'documents. If problems persist, start fresh and recreate the project context manually.'
    )
    doc.add_paragraph()

    doc.add_heading('Need More Help?', 2)
    doc.add_paragraph('Email support: support@pmerit.com')
    doc.add_paragraph('Check for updates in your Gumroad library')
    doc.add_paragraph('Community forum coming soon')
    doc.add_paragraph()
    doc.add_paragraph(
        'When contacting support, include your license email, the AI platform you are using, '
        'a description of the issue, and any relevant error messages or screenshots.'
    )

    doc.add_page_break()


def add_real_world_examples(doc, platform_name=None):
    """Add Real-World Examples chapter"""
    doc.add_heading('Chapter 14: Real-World Examples', 1)

    doc.add_paragraph(
        'Theory is useful, but practice makes perfect. This chapter shows AIXORD in action '
        'across three different project types. Study these examples to see how the framework '
        'adapts to various situations.'
    )
    doc.add_paragraph()

    doc.add_heading('Example 1: Building a Personal Budget Tracker', 2)
    doc.add_paragraph(
        'Sarah wants to build a simple web application to track her monthly expenses. '
        "She is not a programmer but has used AI assistants before. Here is how her "
        'first AIXORD session unfolds.'
    )
    doc.add_paragraph()

    doc.add_heading('Session Start', 3)
    doc.add_paragraph(
        'Sarah pastes the AIXORD governance file into her AI chat and types PMERIT CONTINUE. '
        'After license validation, she states her objective: "Build a personal budget tracker '
        'web app that lets me log expenses, categorize them, and see monthly summaries."'
    )
    doc.add_paragraph()

    doc.add_heading('Discovery Phase', 3)
    doc.add_paragraph(
        'The AI asks clarifying questions: "Who will use this app? Just you, or others? '
        'Do you need it accessible from multiple devices? What expense categories matter to you?" '
        'Sarah answers each question, and the AI documents her responses in the decision ledger.'
    )
    doc.add_paragraph()

    doc.add_heading('Options Phase', 3)
    doc.add_paragraph(
        'The AI presents three approaches: Option A uses a simple spreadsheet with formulas. '
        'Option B uses a no-code tool like Notion. Option C builds a custom web app with React. '
        'Sarah chooses Option B because she wants something more structured than a spreadsheet '
        'but not as complex as custom development.'
    )
    doc.add_paragraph()

    doc.add_heading('Execute Phase', 3)
    doc.add_paragraph(
        'The AI guides Sarah through each step: "Create a new Notion page called '
        'Budget Tracker. Add a database with these columns: Date, Amount, Category, Description." '
        'Sarah follows each instruction and reports back. The AI marks each sub-scope complete '
        'after verification.'
    )
    doc.add_paragraph()

    doc.add_heading('Session End', 3)
    doc.add_paragraph(
        'After 90 minutes, Sarah has a working budget tracker. She types PMERIT HANDOFF and '
        'saves the document. Two days later, she resumes to add a new feature. She pastes the '
        'governance and handoff, types PMERIT CONTINUE, and the AI immediately knows where '
        'she left off.'
    )
    doc.add_paragraph()

    doc.add_heading('Example 2: Refactoring Legacy Code', 2)
    doc.add_paragraph(
        'Marcus is a software developer with a Python codebase that has grown messy over time. '
        'He wants to refactor the authentication module without breaking existing functionality.'
    )
    doc.add_paragraph()

    doc.add_heading('The Objective', 3)
    doc.add_paragraph(
        '"Refactor the authentication module in my Python project. The current code mixes '
        'database access, password hashing, and session management in one file. I want to '
        'separate these concerns while maintaining all existing functionality."'
    )
    doc.add_paragraph()

    doc.add_heading('The Approach', 3)
    doc.add_paragraph(
        'The AI creates a project document with four SCOPEs: SCOPE_AUDIT to analyze existing code '
        'and document current behavior, SCOPE_DESIGN to create new module structure, '
        'SCOPE_REFACTOR to implement the new structure, and SCOPE_VERIFY to run tests and '
        'confirm no regressions. The 4-State Locking system ensures Marcus does not proceed '
        'to refactoring until the audit is verified.'
    )
    doc.add_paragraph()

    doc.add_heading('The Critical Moment', 3)
    doc.add_paragraph(
        'During SCOPE_REFACTOR, Marcus notices the AI suggesting a change that would break '
        'backward compatibility. He issues HALT. The AI stops immediately and awaits direction. '
        'Marcus clarifies that backward compatibility is mandatory. The AI updates the approach '
        'and continues only after Marcus says APPROVED.'
    )
    doc.add_paragraph()

    doc.add_heading('Visual Audit', 3)
    doc.add_paragraph(
        'After refactoring, Marcus provides screenshots of his test output. The AI compares '
        'against the requirements documented in SCOPE_AUDIT: all tests pass, no regressions '
        'detected. The SCOPE_VERIFY state transitions to VERIFIED. Marcus locks the SCOPEs and '
        'his refactoring is complete.'
    )
    doc.add_paragraph()

    doc.add_heading('Example 3: Writing a Business Proposal', 2)
    doc.add_paragraph(
        'Jennifer needs to write a proposal for a new marketing initiative. She uses AIXORD '
        'to structure her thinking and produce a professional document.'
    )
    doc.add_paragraph()

    doc.add_heading('Brainstorm Phase', 3)
    doc.add_paragraph(
        'Jennifer enters BRAINSTORM mode to explore ideas. The AI helps her think through '
        'target audience, budget constraints, timeline, and success metrics. All ideas are '
        'documented but not filtered - this is the creative phase where judgment is suspended.'
    )
    doc.add_paragraph()

    doc.add_heading('Purpose-Bound in Action', 3)
    doc.add_paragraph(
        'Mid-session, Jennifer asks about a competitor product. The AI responds: "This topic '
        'is outside your current project scope. Options: A) Return to proposal, B) Expand scope '
        'to include competitive analysis, C) Save progress and start new session for competitor '
        'research." Jennifer chooses B, expanding scope. The AI incorporates competitive '
        'analysis into the project document.'
    )
    doc.add_paragraph()

    doc.add_heading('Multi-Session Continuity', 3)
    doc.add_paragraph(
        'The proposal takes three sessions across a week. Each session ends with HANDOFF. '
        'Each session resumes with PMERIT CONTINUE. Context is never lost. Decisions made '
        'Monday are still there Friday. The final proposal is comprehensive, well-structured, '
        'and includes everything discussed across all sessions.'
    )
    doc.add_paragraph()

    doc.add_heading('Key Takeaways from These Examples', 2)
    takeaways = [
        'AIXORD adapts to any project type - technical or non-technical',
        'The phase system provides structure without rigidity',
        'HALT is your friend - use it whenever something feels wrong',
        'Handoffs make multi-session projects seamless',
        'Purpose-bound operation keeps focus but allows intentional expansion',
        'The 4-State Locking system prevents regressions',
        'Visual audit applies to any output that can be shown'
    ]
    for t in takeaways:
        doc.add_paragraph('  * ' + t)

    doc.add_page_break()


def add_platform_tips(doc, platform_name, platform_tips):
    """Add Platform-Specific Tips chapter"""
    doc.add_heading(f'Chapter 15: {platform_name} Tips', 1)

    doc.add_paragraph(
        f'This chapter provides tips specifically optimized for {platform_name}. '
        f'While AIXORD works with any AI platform, each has unique characteristics '
        f'that affect how governance should be applied. Use these tips to get the '
        f'most out of your {platform_name} sessions.'
    )
    doc.add_paragraph()

    for heading, content in platform_tips:
        doc.add_heading(heading, 2)
        doc.add_paragraph(content)
        doc.add_paragraph()

    doc.add_heading('General Best Practices', 2)
    doc.add_paragraph(
        'Regardless of platform, these practices improve AIXORD effectiveness: '
        'Always paste the complete governance document at the start of each session. '
        'Use CHECKPOINT regularly to preserve context. End every session with HANDOFF. '
        'Keep sessions focused and time-bounded. When output quality degrades, start fresh.'
    )

    doc.add_page_break()


def add_quick_reference(doc):
    """Add Quick Reference appendix"""
    doc.add_heading('Appendix A: Quick Reference Card', 1)

    doc.add_paragraph(
        'Bookmark this page for quick access during AI sessions. These are the essential '
        'commands and concepts you will use most frequently.'
    )
    doc.add_paragraph()

    doc.add_heading('ACTIVATION', 2)
    doc.add_paragraph('PMERIT CONTINUE - Start a new AIXORD session or resume from handoff')
    doc.add_paragraph()

    doc.add_heading('THE SIX PHASES', 2)
    doc.add_paragraph('DECISION - Default waiting state, AI awaits your direction')
    doc.add_paragraph('DISCOVER - Clarifying unclear ideas through exploration')
    doc.add_paragraph('BRAINSTORM - Generating creative possibilities')
    doc.add_paragraph('OPTIONS - Comparing 2-3 structured alternatives')
    doc.add_paragraph('EXECUTE - Implementing approved plans (requires APPROVED)')
    doc.add_paragraph('AUDIT - Reviewing and verifying completed work')
    doc.add_paragraph()

    doc.add_heading('SESSION CONTROL', 2)
    doc.add_paragraph('APPROVED - Approve current plan and enter execution')
    doc.add_paragraph('HALT - Stop everything, return to decision phase')
    doc.add_paragraph('CHECKPOINT - Save state without ending session')
    doc.add_paragraph('HANDOFF - Generate full export and end session')
    doc.add_paragraph()

    doc.add_heading('PURPOSE-BOUND LEVELS', 2)
    doc.add_paragraph('STRICT - Zero tolerance for off-topic content')
    doc.add_paragraph('STANDARD - Default with polite redirects')
    doc.add_paragraph('RELAXED - Brief tangents allowed')
    doc.add_paragraph('EXPAND SCOPE: [topic] - Add to project boundaries')
    doc.add_paragraph()

    doc.add_heading('ENFORCEMENT', 2)
    doc.add_paragraph('PROTOCOL CHECK - Verify AI compliance with governance')
    doc.add_paragraph('DRIFT WARNING - Flag deviation from rules')
    doc.add_paragraph('VERIFY: [claim] - Request proof of completion')
    doc.add_paragraph()

    doc.add_heading('4-STATE LOCKING', 2)
    doc.add_paragraph('PLANNED - Ready to begin, specification locked')
    doc.add_paragraph('ACTIVE - Under development, changes allowed')
    doc.add_paragraph('IMPLEMENTED - Complete, awaiting audit')
    doc.add_paragraph('VERIFIED - Audited and locked')
    doc.add_paragraph()

    doc.add_heading('KEY PRINCIPLE', 2)
    doc.add_paragraph(
        'You are the Director. AI is your Architect and Commander. Every decision is '
        'documented, every action is authorized, nothing is forgotten between sessions.'
    )

    doc.add_page_break()


def add_glossary(doc):
    """Add Glossary appendix"""
    doc.add_heading('Appendix B: Glossary of Terms', 1)

    terms = [
        ('AIXORD', 'AI Execution Order - the governance framework for structured AI collaboration.'),
        ('Approval Gate', 'A checkpoint requiring explicit human APPROVED command before AI executes.'),
        ('Architect', "The AI's planning role - analyzes, recommends, and specifies but does not implement."),
        ('Behavioral Firewall', 'Rules that suppress unwanted AI default behaviors like verbosity.'),
        ('Carryforward', 'Information explicitly marked to persist across sessions via handoff.'),
        ('Checkpoint', 'Mid-session state save that preserves context without ending the session.'),
        ('Choice Elimination', 'Rule preventing AI from offering multiple unsolicited options.'),
        ('Commander', "The AI's execution role - implements approved plans and delivers artifacts."),
        ('Default Suppression', "Rule keeping AI's default state minimal - only what is requested."),
        ('Director', "The human's role - makes all decisions, provides approvals, owns outcomes."),
        ('Drift', 'When AI gradually deviates from governance compliance or project focus.'),
        ('Execute Phase', 'The implementation phase requiring explicit APPROVED to enter.'),
        ('Expansion Trigger', 'Keywords like EXPLAIN or WHY that permit verbose AI output.'),
        ('Governance', 'The complete set of rules and protocols controlling AI behavior.'),
        ('Handoff', 'Full session state export document for continuing work in new sessions.'),
        ('Hard Stop', 'Rule requiring AI to stop completely after task completion.'),
        ('Instruction Priority', 'Hierarchy: Governance > User Commands > Task Content > Training.'),
        ('Phase', 'One of six operational modes that define allowed AI behaviors.'),
        ('Purpose-Bound', 'Principle that AI operates exclusively within stated project objective.'),
        ('Reasoning Trace', "Visible step-by-step explanation of AI's decision process."),
        ('Redirect Protocol', 'Steps AI follows when handling out-of-scope requests.'),
        ('Response Header', 'Mandatory status display showing phase and progress on every response.'),
        ('SCOPE', 'A documented deliverable unit within the project hierarchy.'),
        ('State File', 'JSON document tracking session configuration, decisions, and progress.'),
        ('SUB-SCOPE', 'A smaller unit of work within a SCOPE.'),
        ('4-State Locking', 'Status progression: PLANNED -> ACTIVE -> IMPLEMENTED -> VERIFIED.'),
        ('Visual Audit', 'Verification process comparing screenshots against specifications.')
    ]

    for term, definition in terms:
        p = doc.add_paragraph()
        p.add_run(term + ': ').bold = True
        p.add_run(definition)

    doc.add_page_break()


def add_notes_pages(doc):
    """Add Notes pages"""
    doc.add_heading('Notes', 1)

    doc.add_paragraph(
        'Use these pages to record your AIXORD customizations, frequently-used commands, '
        'and insights from your sessions. Having a written record helps you refine your '
        'approach over time.'
    )
    doc.add_paragraph()

    doc.add_heading('My Project Objectives', 2)
    doc.add_paragraph('Record the objectives you use most frequently:')
    for _ in range(6):
        doc.add_paragraph('_' * 55)
    doc.add_paragraph()

    doc.add_heading('My Preferred Settings', 2)
    doc.add_paragraph('Purpose-Bound Level: ' + '_' * 35)
    doc.add_paragraph('Preferred AI Platform: ' + '_' * 35)
    doc.add_paragraph('Default Phase for New Projects: ' + '_' * 25)
    doc.add_paragraph('Checkpoint Frequency: ' + '_' * 35)
    doc.add_paragraph()

    doc.add_heading('Commands I Use Most', 2)
    doc.add_paragraph('Track which commands you use frequently:')
    for _ in range(6):
        doc.add_paragraph('_' * 55)
    doc.add_paragraph()

    doc.add_heading('Custom Scope Boundaries', 2)
    doc.add_paragraph('Document scope boundaries that work well for your projects:')
    for _ in range(4):
        doc.add_paragraph('_' * 55)

    doc.add_page_break()

    doc.add_heading('Lessons Learned', 1)
    doc.add_paragraph(
        'After each significant project, record what worked well and what you would do '
        'differently next time. These lessons become invaluable over time.'
    )
    doc.add_paragraph()

    for _ in range(20):
        doc.add_paragraph('_' * 55)

    doc.add_page_break()

    # Session Log page
    doc.add_heading('Session Log', 1)
    doc.add_paragraph(
        'Track your AIXORD sessions here. Recording dates, platforms, and outcomes '
        'helps you identify patterns and improve your workflow.'
    )
    doc.add_paragraph()

    for i in range(1, 11):
        doc.add_paragraph(f'Session {i}')
        doc.add_paragraph('Date: _______ Platform: _______ Project: _______')
        doc.add_paragraph('Key Outcome: ' + '_' * 40)
        doc.add_paragraph()

    doc.add_page_break()

    # Third notes page for additional space
    doc.add_heading('Additional Notes', 1)
    doc.add_paragraph()

    for _ in range(25):
        doc.add_paragraph('_' * 55)

    doc.add_page_break()


def add_final_page(doc, title):
    """Add final page"""
    for _ in range(3):
        doc.add_paragraph()

    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.add_run('AIXORD v3.2.1').bold = True

    doc.add_paragraph()

    p = doc.add_paragraph('Purpose-Bound. Disciplined. Focused.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph('Transform chaotic AI conversations into structured, productive projects.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph('Copyright 2026 PMERIT LLC. All Rights Reserved.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph('For companion templates and digital files visit:')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph('pmerit.gumroad.com')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph()

    p = doc.add_paragraph('Questions or support: support@pmerit.com')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph('Thank you for choosing AIXORD.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph('Now go build something amazing.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True


def count_words(doc):
    """Count words in document"""
    return sum(len(p.text.split()) for p in doc.paragraphs)


# Platform-specific configurations
MANUSCRIPTS = {
    'AIXORD_FOR_DEEPSEEK_USERS.docx': {
        'title': 'AIXORD for DeepSeek Users',
        'subtitle': 'AI Governance Optimized for DeepSeek',
        'platform': 'DeepSeek',
        'url': 'chat.deepseek.com',
        'tips': [
            ('Leverage Deep Reasoning',
             'DeepSeek excels at step-by-step reasoning. For complex problems, ask for detailed '
             'reasoning traces. The model will show its work, making it easier to verify logic. '
             'Use phrases like "think through this step by step" to activate deep reasoning mode.'),
            ('Code Generation Excellence',
             'DeepSeek Coder is optimized for programming. Request complete, production-ready '
             'implementations with inline documentation. The model handles complex codebases well '
             'and can maintain context across multiple files.'),
            ('Knowledge Recency Awareness',
             'DeepSeek training data has a cutoff date. Always verify current versions of APIs, '
             'libraries, and pricing. When working with external services, ask the model to note '
             'which information may need verification.'),
            ('Mathematical Precision',
             'DeepSeek handles complex calculations well. Ask for intermediate steps to verify '
             'accuracy. The model can show its mathematical reasoning clearly.'),
            ('Cost Efficiency',
             'DeepSeek offers competitive pricing. For large projects, this allows more iterations '
             'and experimentation without budget concerns.'),
            ('Context Window Usage',
             'DeepSeek handles long contexts well. You can paste full governance documents plus '
             'handoffs plus project files without truncation issues.')
        ]
    },
    'AIXORD_FOR_UNIVERSAL_USERS.docx': {
        'title': 'AIXORD Universal',
        'subtitle': 'AI Governance for Any Platform',
        'platform': 'Any AI',
        'url': 'your preferred AI chat',
        'tips': [
            ('Cross-Model Compatibility',
             'AIXORD Universal works with any AI platform: Claude, ChatGPT, Gemini, DeepSeek, '
             'Copilot, and more. The governance document adapts its behavior based on which '
             'model is responding.'),
            ('Model Self-Identification',
             'Use the IDENTIFY MODEL command to see which AI is currently responding and what '
             'capabilities it has detected. This helps you understand the current context.'),
            ('Seamless Model Switching',
             'Start a project in ChatGPT, continue in Claude, finish in DeepSeek. The handoff '
             'format is standardized across all platforms, ensuring smooth transitions.'),
            ('Capability Assessment',
             'Different models have different strengths. AIXORD Universal detects capabilities '
             'like file uploads, code execution, and image analysis, adjusting its behavior accordingly.'),
            ('Portability',
             'Save governance files and handoffs locally. You can resume your project with any '
             'AI that can read text. No vendor lock-in.'),
            ('Future-Proof',
             'New AI models released in the future will work with AIXORD Universal. The framework '
             'is based on natural language instructions that any capable model can follow.')
        ]
    },
    'AIXORD_FOR_GEMINI_USERS.docx': {
        'title': 'AIXORD for Gemini Users',
        'subtitle': 'AI Governance Optimized for Google Gemini',
        'platform': 'Gemini',
        'url': 'gemini.google.com',
        'tips': [
            ('Google Workspace Integration',
             'Gemini integrates well with Google Docs, Sheets, and other Google services. '
             'Use this for projects that involve Google Workspace products.'),
            ('Research and Factual Queries',
             'Gemini handles factual queries and research tasks effectively. It can access '
             'current information through Google Search integration.'),
            ('Google Colab for Code',
             'For code-heavy projects, combine Gemini with Google Colab for execution. '
             'This gives you a complete development environment.'),
            ('Context Length',
             'Gemini Advanced provides extended context length needed for full governance '
             'documents. Free tier may require abbreviated governance.'),
            ('Multimodal Capabilities',
             'Gemini can analyze images. Use this for visual audit of UI implementations. '
             'Upload screenshots directly for verification.'),
            ('Version Awareness',
             'Gemini Pro and Gemini Ultra have different capabilities. Ensure you are using '
             'the appropriate version for your project complexity.')
        ]
    },
    'AIXORD_FOR_COPILOT_USERS.docx': {
        'title': 'AIXORD for Copilot Users',
        'subtitle': 'AI Governance Optimized for Microsoft Copilot',
        'platform': 'Microsoft Copilot',
        'url': 'copilot.microsoft.com',
        'tips': [
            ('Microsoft 365 Integration',
             'Copilot integrates natively with Word, Excel, PowerPoint, and Outlook. '
             'Use it for tasks involving Microsoft Office products.'),
            ('Enterprise Features',
             'Enterprise Copilot maintains conversation history across sessions, reducing '
             'the need for handoff documents in some cases.'),
            ('Bing Search Access',
             'Copilot can access Bing search for up-to-date information retrieval. '
             'This helps verify current data and pricing.'),
            ('GPT-4 Access',
             'For best results, use Copilot Pro with GPT-4 access enabled. This provides '
             'stronger instruction-following capabilities.'),
            ('Code Assistance',
             'Copilot in Visual Studio Code provides inline code suggestions. AIXORD governance '
             'can guide how you use these suggestions in your projects.'),
            ('Teams Integration',
             'Copilot in Microsoft Teams can help with meeting summaries and action items. '
             'Use AIXORD structure to organize these outputs.')
        ]
    },
    'AIXORD_FOR_CLAUDE_USERS.docx': {
        'title': 'AIXORD for Claude Users',
        'subtitle': 'AI Governance Optimized for Anthropic Claude',
        'platform': 'Claude',
        'url': 'claude.ai',
        'tips': [
            ('Instruction Following',
             'Claude excels at following complex, nuanced instructions. The full AIXORD governance '
             'document works particularly well with Claude, as it carefully follows each rule.'),
            ('Long Context',
             'Claude handles very long contexts well. You can paste full governance documents, '
             'handoffs, and extensive project files without losing coherence.'),
            ('Writing Quality',
             'Claude produces high-quality written content. For documentation, proposals, and '
             'creative writing projects, leverage this strength.'),
            ('Code Explanation',
             'Claude is excellent at explaining code. Use the EXPLAIN trigger word to get '
             'detailed explanations of complex implementations.'),
            ('Artifacts Feature',
             'Claude can create standalone artifacts (documents, code files). Use this for '
             'deliverables that need to be saved separately.'),
            ('Projects Feature',
             'Claude Projects allow persistent context across conversations. This can supplement '
             'HANDOFF documents for ongoing work.')
        ]
    },
    'AIXORD_FOR_CHATGPT_USERS.docx': {
        'title': 'AIXORD for ChatGPT Users',
        'subtitle': 'AI Governance Optimized for OpenAI ChatGPT',
        'platform': 'ChatGPT',
        'url': 'chat.openai.com',
        'tips': [
            ('Model Selection',
             'GPT-4 follows AIXORD governance more reliably than GPT-3.5. Use ChatGPT Plus '
             'or Team for the best experience with full governance.'),
            ('Custom GPTs',
             'You can create a Custom GPT with AIXORD governance pre-loaded. This eliminates '
             'the need to paste governance at the start of each session.'),
            ('Code Interpreter',
             'ChatGPT Code Interpreter can execute Python code. Use this for data analysis '
             'projects where you need to run and verify code.'),
            ('DALL-E Integration',
             'ChatGPT can generate images with DALL-E. For projects involving visual assets, '
             'this can be incorporated into your AIXORD workflow.'),
            ('Browsing Capability',
             'ChatGPT Plus can browse the web for current information. Use this to verify '
             'dates, pricing, and API versions.'),
            ('Memory Feature',
             'ChatGPT memory can persist some context between sessions. This supplements '
             'but does not replace HANDOFF documents.')
        ]
    },
    'AIXORD_GENESIS.docx': {
        'title': 'AIXORD Genesis',
        'subtitle': 'The Origin and Philosophy of AI Governance',
        'platform': None,
        'url': None,
        'tips': [
            ('The Genesis Philosophy',
             'AIXORD Genesis represents the foundational thinking behind AI governance. '
             'Understanding the philosophy helps you apply the framework more effectively.'),
            ('Why Structure Matters',
             'AI systems are capable but directionless without structure. Genesis explains '
             'why governance transforms AI from a toy into a tool.'),
            ('The Authority Principle',
             'Humans must remain in control of AI systems. Genesis establishes why this '
             'principle is non-negotiable for productive collaboration.'),
            ('Evolution of the Framework',
             'AIXORD has evolved through real-world use. Genesis documents the lessons '
             'learned and how they shaped the current framework.'),
            ('Future Directions',
             'As AI capabilities grow, governance must evolve. Genesis provides the '
             'philosophical foundation for future AIXORD development.'),
            ('Community Contribution',
             'AIXORD Genesis invites users to contribute insights. Your experiences '
             'help improve the framework for everyone.')
        ]
    },
    'AIXORD_BUILDERS_TOOLKIT.docx': {
        'title': 'AIXORD Builders Toolkit',
        'subtitle': 'Tools and Templates for Project Success',
        'platform': None,
        'url': None,
        'tips': [
            ('Template Library',
             'The Builders Toolkit includes templates for common project types: web apps, '
             'APIs, documentation, marketing campaigns, and more.'),
            ('SCOPE Templates',
             'Pre-built SCOPE templates help you decompose projects quickly. Customize '
             'them for your specific needs.'),
            ('Handoff Templates',
             'Standardized handoff templates ensure consistent session transfers. '
             'No important context is ever lost.'),
            ('Audit Checklists',
             'Pre-built audit checklists for visual verification. Ensure nothing '
             'is missed during the AUDIT phase.'),
            ('Command Quick Reference',
             'Printable command reference cards for your desk. Quick access to '
             'all AIXORD commands during sessions.'),
            ('Project Starter Kits',
             'Complete starter kits for common project types. Begin with proven '
             'structures rather than starting from scratch.')
        ]
    },
    'AIXORD_THE_COMPLETE_FRAMEWORK.docx': {
        'title': 'AIXORD: The Complete Framework',
        'subtitle': 'Comprehensive AI Governance Reference',
        'platform': None,
        'url': None,
        'tips': [
            ('Comprehensive Coverage',
             'The Complete Framework documents every AIXORD feature in detail. '
             'Use this as your authoritative reference.'),
            ('Advanced Techniques',
             'Learn advanced governance techniques: multi-project management, '
             'team collaboration, and enterprise deployment.'),
            ('Edge Cases',
             'The Complete Framework covers edge cases and unusual situations. '
             'Know how to handle anything that arises.'),
            ('Customization Guide',
             'Learn how to customize AIXORD for your specific needs. Modify '
             'enforcement levels, add custom commands, and more.'),
            ('Integration Patterns',
             'Patterns for integrating AIXORD with other tools and workflows. '
             'Make governance part of your complete system.'),
            ('Troubleshooting Deep Dive',
             'Detailed troubleshooting for complex issues. When basic fixes '
             'do not work, consult this section.')
        ]
    },
    'AIXORD_STARTER_GUIDE.docx': {
        'title': 'AIXORD Starter Guide',
        'subtitle': 'Your First Steps with AI Governance',
        'platform': None,
        'url': None,
        'tips': [
            ('Quick Start Focus',
             'The Starter Guide gets you productive in 15 minutes. It covers '
             'only the essentials needed to begin.'),
            ('Learning Path',
             'Follow the suggested learning path: Quick Start, then Core Concepts, '
             'then Reference sections as needed.'),
            ('Practice Projects',
             'Use the included practice projects to build skills before tackling '
             'your real work.'),
            ('Common Pitfalls',
             'The Starter Guide highlights the most common beginner mistakes '
             'so you can avoid them from the start.'),
            ('Upgrade Path',
             'When ready for more, the Starter Guide points you to advanced '
             'resources in the AIXORD ecosystem.'),
            ('Support Resources',
             'Contact support@pmerit.com for help getting started. We want '
             'your first AIXORD experience to be successful.')
        ]
    }
}


def expand_manuscript(filename, config, output_dir):
    """Expand a single manuscript to 24+ pages"""
    doc = Document()

    # Add all sections
    add_front_matter(doc, config['title'], config['subtitle'], config.get('platform'))
    add_preface(doc, config.get('platform'))
    add_introduction(doc, config.get('platform'))
    add_what_is_aixord(doc)
    add_quick_start(doc, config.get('platform'), config.get('url'))
    add_authority_model(doc)
    add_phase_system(doc)
    add_session_continuity(doc)
    add_purpose_bound(doc)
    add_behavioral_firewalls(doc)
    add_project_decomposition(doc)
    add_locking_system(doc)
    add_visual_audit(doc)
    add_common_mistakes(doc)
    add_commands_reference(doc)
    add_troubleshooting(doc, config.get('platform'))
    add_real_world_examples(doc, config.get('platform'))

    if config.get('tips'):
        add_platform_tips(doc, config.get('platform') or 'This Edition', config['tips'])

    add_quick_reference(doc)
    add_glossary(doc)
    add_notes_pages(doc)
    add_final_page(doc, config['title'])

    # Save
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    # Report
    words = count_words(doc)
    pages = words / 275
    status = 'OK' if pages >= 24 else 'NEEDS MORE'
    print(f'{filename}: {words} words, {pages:.1f} pages - {status}')

    return words, pages


def main():
    output_dir = r'C:\dev\pmerit\Pmerit_Product_Development\products\aixord-chatbot\manuscripts\kdp'

    print('Expanding all manuscripts to 24+ pages...')
    print('=' * 60)

    results = []
    for filename, config in MANUSCRIPTS.items():
        words, pages = expand_manuscript(filename, config, output_dir)
        results.append((filename, words, pages))

    print('=' * 60)
    print('Summary:')
    all_ok = True
    for filename, words, pages in results:
        status = 'OK' if pages >= 24 else 'NEEDS MORE'
        if pages < 24:
            all_ok = False
        print(f'  {filename}: {pages:.1f} pages - {status}')

    if all_ok:
        print('\nAll manuscripts meet 24-page requirement!')
    else:
        print('\nSome manuscripts need additional content.')


if __name__ == '__main__':
    main()
